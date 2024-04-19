from tkinter import*
from customtkinter import*
import customtkinter as ctk
from PIL import Image
from tkinter import messagebox
import sqlite3 
import time
import docx
from docx import Document
import permission
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import win32api
import win32print
import os
import sys

def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS2
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)

ctk.set_appearance_mode("system")
ctk.set_default_color_theme("blue")


class MyFrame(ctk.CTkFrame):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        

class ConsultationClass(ctk.CTk):
    def __init__(self,root):
        self.root=root
        self.root.state('zoomed')
        # Bind the exit method to the window's close button
        self.root.protocol("WM_DELETE_WINDOW", self.exit)
        screen_width=self.root.winfo_screenwidth()
        screen_height=self.root.winfo_screenheight()
        self.root.geometry(f'{screen_width}x{screen_height}')
        self.root.iconbitmap(resource_path('icon.ico'))
        self.root.title("SmileScribePro")
        self.chv=0
        self.chk_print=0

        #============================Title======================
        self.icon_title=ctk.CTkImage(dark_image=Image.open(resource_path("images\\logo.png")),
                                    light_image=Image.open(resource_path("images\\logo.png")),size=(50,50))
        title = ctk.CTkLabel(
        self.root,
        text="SmileScribePro",
        image=self.icon_title,
        compound=LEFT,
        font=("sans-serif", 40, "bold"),
        fg_color=("#3498db","#fff"),
        text_color=("#fff","#3498db"),
        anchor='w',
        width=screen_width,  # Set the width to be the screen width
        height=68,padx=20
    ).place(anchor='nw', x=0, y=0)
        #=============Btn_Logout================================
        # Calculate the position of the button based on the screen size
        x_position = screen_width - 100  # Adjust this value as needed
        y_position = screen_height - 735  # Adjust this value as needed
        btn_logout = ctk.CTkButton(self.root,text="Exit"
                                    ,command=self.exit
                                    ,font=("calibri",20,"bold")
                                    ,fg_color=("#273c75","#353b48")
                                    ,hover_color="#e84118"
                                    ,height=50,width=150
                                    ,corner_radius=0
                                    ).place(anchor='e', x=x_position, y=30)
        #======================Clock============================
        self.lbl_clock=ctk.CTkLabel(self.root,text="Welcome to SmileScribePro - Professional Patient Records Management System\t\t Date: DD-MM-YYYY\t\t Time: HH:MM:SS"
                                    ,font=("calibri",18,"bold")
                                    ,fg_color=("#487eb0","#4d636d"),bg_color="#fff"
                                    ,width=screen_width
                                    ,text_color="#fff")
        self.lbl_clock.place(x=0,y=65)

        #===================title=========
        label_width = screen_width - 40
        title=ctk.CTkLabel(self.root,text="Generate Patient Consultation Sheet",font=("helvetica new",20,"bold"),fg_color=("#0f4d7d","#fff"),text_color=("white","#000"),height=40,width=label_width,padx=20).place(x=20,y=93)
        
        
        frame1_y=screen_height-620
        frame1_x=screen_width-1350
        consultframe=MyFrame(self.root)
        consultframe.place(x=10,y=150)
        title=ctk.CTkLabel(consultframe,text="Enter Patient Details",font=("times new roman",18,"bold")).grid(row=0, column=0, sticky='ew')
        #================content================
        #==========Variables============
        self.var_pname=ctk.StringVar()
        self.var_add=ctk.StringVar()
        self.var_phone=ctk.StringVar()
        self.var_prof=ctk.StringVar()
        self.var_age=ctk.StringVar()
        
        lbl_patient_name=ctk.CTkLabel(consultframe,text="Patient Names",font=("goudy old style",18),padx=40,pady=20).grid(row=2,column=0,sticky='w')
        lbl_address=ctk.CTkLabel(consultframe,text="Address",font=("goudy old style",18),padx=40,pady=20).grid(row=3,column=0,sticky='w')
        lbl_phone=ctk.CTkLabel(consultframe,text="Phone No.",font=("goudy old style",18),padx=40,pady=20).grid(row=4,column=0,sticky='w')
        lbl_profession=ctk.CTkLabel(consultframe,text="Profession",font=("goudy old style",18),padx=40,pady=20).grid(row=5,column=0,sticky='w')
        lbl_age=ctk.CTkLabel(consultframe,text="Age",font=("goudy old style",18),padx=40,pady=20).grid(row=6,column=0,sticky='w')
        
        
        txt_patient_name=ctk.CTkEntry(consultframe,textvariable=self.var_pname,font=("goudy old style",18),fg_color="lightyellow").place(x=160,y=50)
        txt_address=ctk.CTkEntry(consultframe,textvariable=self.var_add,font=("goudy old style",18),fg_color="lightyellow").place(x=150,y=110)
        txt_phone=ctk.CTkEntry(consultframe,textvariable=self.var_phone,font=("goudy old style",18),fg_color="lightyellow").place(x=150,y=170)
        txt_profession=ctk.CTkEntry(consultframe,textvariable=self.var_prof,font=("goudy old style",15),fg_color="lightyellow").place(x=150,y=240)
        txt_age=ctk.CTkEntry(consultframe,textvariable=self.var_age,font=("goudy old style",18),fg_color="lightyellow").place(x=150,y=300)
        
        #==================Button===========
        
        btn_generate=ctk.CTkButton(consultframe,command=self.generate_file,text="Generate Sheet",font=("times new roman",15),fg_color="#44bd32",hover_color="#badc58",width=150).grid(row=8,column=0)
        btn_print=ctk.CTkButton(consultframe,command=self.print_file,text="Print Sheet",font=("times new roman",15),fg_color="#eb2f06",width=150,hover_color="#f0932b").grid(row=8,column=1)
        btn_print=ctk.CTkButton(consultframe,command=self.view_sheet,text="View Sheet",font=("times new roman",15),fg_color="#2c3e50",width=150,hover_color="#9b59b6").grid(row=9,column=0)
        btn_clear=ctk.CTkButton(consultframe,command=self.clear,text="Clear All",font=("times new roman",15),fg_color="#3c6382",hover_color="#4834d4",width=150).grid(row=9,column=1)
        
        #==================File Area===========
        frame2_x=screen_width-950
        frame2_y=screen_height-620
        fileFrame=MyFrame(self.root,width=screen_width,height=screen_height)
        fileFrame.place(x=400,y=150)
        
        BTitle=ctk.CTkLabel(fileFrame,text="View Patient Consultation Sheet Area",font=("goudy old style",20,"bold")).pack(side=TOP,fill=X)
        scrolly=Scrollbar(fileFrame,orient=VERTICAL)
        scrollx=Scrollbar(fileFrame,orient=HORIZONTAL)
        scrolly.pack(side=RIGHT,fill=Y)
        scrollx.pack(side=BOTTOM,fill=X)
        
        self.txt_file_area=ctk.CTkTextbox(fileFrame,font=("Courier",18,"bold"),width=int(screen_width-450),height=int(screen_height-250),wrap=WORD,yscrollcommand=scrolly.set,xscrollcommand=scrollx.set)
        self.txt_file_area.pack(fill=BOTH,expand=1)
        scrolly.config(command=self.txt_file_area.yview)
        scrollx.config(command=self.txt_file_area.xview)
        
        #==================Footer================
        date_=time.strftime("%Y")
        lbl_footer=ctk.CTkLabel(self.root,text=f"Copyright @ {date_} RootTech", font=("times new roman",12,"bold")).pack(side=BOTTOM,fill=X)
       
        permission.interact_with_database((resource_path('PRMS.db')))  
        self.update_date_time()


        #=========================All Functions==========
    
    def generate_file(self):
        if self.var_pname.get()=='' or self.var_add.get()=='' or self.var_phone.get()=="" or self.var_prof.get()=="" or self.var_age.get()=="":
            messagebox.showerror("Error",f"All Patient Details are required",parent=self.root)
        else:
            #======file TOP=====
            self.file_top()
            messagebox.showinfo('Congrats',"file has been generated",parent=self.root)
            self.chk_print=1
            self.chk_view=1
    
    def find_or_create_table(self,doc):
                for element in doc.element.body:
                    if isinstance(element, docx.oxml.table.CT_Tbl):
                        return docx.table.Table(element, doc)  # Use the existing table

                    # Create a new table with headers
                table = doc.add_table(rows=6, cols=6)
                table.style = 'Table Grid'
                table.autofit = True
                table.allow_autofit = True
                header_row = table.rows[0]
                header_row.cells[0].text = 'DATE'
                header_row.cells[1].text = 'TOOTH'
                header_row.cells[2].text = 'NATURE OF INTERVENTION '
                header_row.cells[3].text = 'AMOUNT DUE'
                header_row.cells[4].text = 'AMOUNT PAID'
                header_row.cells[5].text = 'BALANCE'

                return table
    
    def file_top(self):
        doc=docx.Document()
        header=doc.sections[0].header
        htable=header.add_table(1,2,Inches(6))
        htab_cells=htable.rows[0].cells
        
        ht0=htab_cells[0].add_paragraph()
        logo=ht0.add_run()
        logo.add_picture('images\\emma.png')
        
        ht1=htab_cells[1].add_paragraph(f'''\tEMMANUEL DENTAL CLINIC 
        BONAMOUSSADI
        Pr.(Dr) AGBOR MICHEAL ASHU
        Tel: 677 17 01 67/697 12 27 82
        www.emmanueldentalcare.org
        ''')
        ht1.alignment=WD_ALIGN_PARAGRAPH.CENTER
        
        title=doc.add_heading(f'''PATIENT CONSULTATION SHEET
''',level=0)
        title.alignment=WD_ALIGN_PARAGRAPH.CENTER
        
        #patient Details head
        doc.add_heading('PATIENT DETAILS')
        #patient Details
        details=doc.add_paragraph()
        details.add_run('Full Name: ')
        details.add_run(f"{self.var_pname.get()}")
        details.add_run('\n')
        details.add_run('Address: ')
        details.add_run(f"{self.var_add.get()}")
        details.add_run('\n')
        details.add_run('Phone: ')
        details.add_run(f"{self.var_phone.get()}")
        details.add_run('\n')
        details.add_run('Age: ')
        details.add_run(f"{self.var_age.get()}")
        details.add_run('\n')
        
        #observations
        
        doc.add_heading('OBSERVATIONS')
        ob=doc.add_paragraph()
        ob.add_run('Doctor: ')
        ob.add_run(f'......................................................................................................')
        ob.add_run('\n')
        ob.add_run('Main Complain: ')
        ob.add_run(f'..........................................................................................................................................................................')
        ob.add_run(f'....................................................................................................................................................................................................................')
        ob.add_run('\n')
        ob.add_run('Treatment Plan: ')
        ob.add_run(f'..........................................................................................................................................................................')
        ob.add_run(f'....................................................................................................................................................................................................................')
        ob.add_run('\n')
        ob.add_run('Observations: ')
        ob.add_run(f'..........................................................................................................................................................................')
        ob.add_run(f'....................................................................................................................................................................................................................')
        #table
        d=time.strftime("%d/%m/%Y")
            # Create a new table with headers
        table = self.find_or_create_table(doc)

        # Add a new row with data
        new_row = table.add_row().cells
        new_row[0].text = ''  #date
        new_row[1].text = ''  # Tooth
        new_row[2].text = ''  # Nature of intervention
        new_row[3].text = ''  # Due
        new_row[4].text = ''  # Paid
        new_row[5].text = ''  # Balance
        
        doc.save(resource_path('consultation\\consultation_sheet.docx'))
        file_path=resource_path('consultation\\consultation_sheet.docx')
        docx_text=self.read_docx(file_path)
        self.txt_file_area.insert('1.0',docx_text)
        

    def read_docx(self,file_path):
        doc=docx.Document(resource_path(file_path))
        fulltext=[]
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                fulltext.append(paragraph.text)
        for para in doc.paragraphs:
            fulltext.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    fulltext.append(cell.text)
        return '\n'.join(fulltext)
        
    def view_sheet(self):
        if self.chk_view==1:
            os.startfile(os.path.join(os.getcwd(),resource_path(f'consultation\\consultation_sheet.docx')))         
        else:
            messagebox.showerror("View Error","Generate Patient sheet First before viewing",parent=self.root)
     
        
    def print_file(self):
        if self.chk_print==1:
            file_path=resource_path('consultation\\consultation_sheet.docx')
            doc=docx.Document(file_path)
            printer_name=win32print.GetDefaultPrinter()
            win32api.ShellExecute(0,"print",file_path,'d:"%s"'%printer_name,".",0)
            self.add()
        else:
            messagebox.showerror("View Error","Generate Patient sheet First before printing",parent=self.root)
        
    def clear(self):
        self.var_add.set("")
        self.var_age.set("")
        self.var_phone.set("")
        self.var_pname.set("")
        self.var_prof.set("")        
        self.txt_file_area.insert("1.0",END)
        self.txt_file_area.delete("1.0",END)

    def update_date_time(self):   
        time_=time.strftime("%H:%M:%S")
        date_=time.strftime("%d:%m:%Y")
        self.lbl_clock.configure(text=f"Welcome to SmileScribePro - Professional Patients Record Management System\t\t Date: {str(date_)}\t\t Time: {str(time_)}")
        self.lbl_clock.after(200,self.update_date_time)
    
    def add(self):
        permission.interact_with_database((resource_path('PRMS.db')))
        con=sqlite3.connect(database=resource_path(r'PRMS.db'))
        cur=con.cursor()
        try:
            if self.var_pname.get()=="":
                messagebox.showerror("Error","Patient Name is required",parent=self.root)
            else:
                cur.execute("Select * from patient where name=?",(self.var_pname.get(),))
                row=cur.fetchone()
                if row!=None:
                    messagebox.showerror("Error","This Patient is already registered, try a different",parent=self.root)
                else:
                    cur.execute("Insert into patient (name,address,phone,profession,dob) values(?,?,?,?,?)",(
                                                self.var_pname.get(),                                                
                                                self.var_add.get(),
                                                self.var_phone.get(),
                                                self.var_prof.get(),                                                
                                                self.var_age.get(),         
                        
                    ))
                    con.commit()
                    messagebox.showinfo("Success","Patient Record Added Successfully",parent=self.root)
                    self.clear()
        except Exception as ex:
            messagebox.showerror("Error",f"Error due to : {str(ex)}",parent=self.root)

    
      
    def exit(self):
        self.is_running = False
        self.root.destroy() 
        
        
            
    
if __name__=="__main__":
    root=ctk.CTk()
    obj=ConsultationClass(root)
    root.mainloop()