from tkinter import*
import customtkinter as ctk
from PIL import Image
import tkinter as tk
from tkinter import ttk,messagebox
import sqlite3
import time
from datetime import datetime
import docx
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import win32print
import win32api
import os
import permission
import sys
from plyer import notification


def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS2
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)



ctk.set_appearance_mode("system")
ctk.set_default_color_theme("blue")


class PlaceholderEntry(ctk.CTkEntry):
    def __init__(self, master=None, placeholder_text="PLACEHOLDER", *args, **kwargs):
        super().__init__(master, *args, **kwargs)
        self.placeholder = placeholder_text

    def get(self):
        # If the Entry is empty, return the placeholder
        if super().get() == "":
            return self.placeholder
        else:
            return super().get()




class archiveClass(ctk.CTk):
    def __init__(self,root):
        self.root=root
        self.root.state('zoomed')
        # Bind the exit method to the window's close button
        self.root.protocol("WM_DELETE_WINDOW", self.exit)
        screen_width=self.root.winfo_screenwidth()
        screen_height=self.root.winfo_screenheight()
        self.root.geometry(f'{screen_width}x{screen_height}')
        self.root.iconbitmap(resource_path('icon.ico'))
        self.root.title("SmileScribePro")
        #=============All Variables=======
        self.var_searchtxt=ctk.StringVar()        
        
        self.var_pat_id=ctk.StringVar()
        self.var_name=ctk.StringVar()
        self.var_address=ctk.StringVar()
        self.var_phone=ctk.StringVar()
        self.var_prof=ctk.StringVar()
        self.var_dob=ctk.StringVar()
        self.var_gender=ctk.StringVar()
        self.var_mc=ctk.StringVar()
        self.var_tp=ctk.StringVar()
        self.var_tooth=ctk.StringVar()
        self.var_doctor=ctk.StringVar()
        
       
       
        self.doc_list=[]
        self.fetch_doctor()
               
        #============================Title======================
        self.icon_title=ctk.CTkImage(dark_image=Image.open(resource_path("images\\logo.png")),
                                    light_image=Image.open(resource_path("images\\logo.png")),size=(50,50))
        title = ctk.CTkLabel(
        self.root,
        text="SmileScribePro",
        image=self.icon_title,
        compound=LEFT,
        font=("sans-serif", 40, "bold"),
        fg_color=("#3498db","#fff"),
        text_color=("#fff","#3498db"),
        anchor='w',
        width=screen_width,  # Set the width to be the screen width
        height=68,padx=20
    ).place(anchor='nw', x=0, y=0)
        #=============Btn_Logout================================
        # Calculate the position of the button based on the screen size
        x_position = screen_width - 100  # Adjust this value as needed
        y_position = screen_height - 735  # Adjust this value as needed
        btn_logout = ctk.CTkButton(self.root,text="Exit"
                                    ,command=self.exit
                                    ,font=("calibri",20,"bold")
                                    ,fg_color=("#273c75","#353b48")
                                    ,hover_color="#e84118"
                                    ,height=50,width=150
                                    ,corner_radius=0
                                    ).place(anchor='e', x=x_position, y=30)
        #======================Clock============================
        self.lbl_clock=ctk.CTkLabel(self.root,text="Welcome to SmileScribePro - Professional Patient Records Management System\t\t Date: DD-MM-YYYY\t\t Time: HH:MM:SS"
                                    ,font=("calibri",18,"bold")
                                    ,fg_color=("#487eb0","#4d636d"),bg_color="#fff"
                                    ,width=screen_width
                                    ,text_color="#fff")
        self.lbl_clock.place(x=0,y=65)

        
        #===================title=========
        title=ctk.CTkLabel(self.root,text="ARCHIVES",font=("Cambria",22,"bold"),fg_color="#0f4d7d",text_color="#fff",width=1000).place(x=170,y=110)
       
        #================search Frame=============
        #==========search==============
        lbl_serach=ctk.CTkLabel(self.root,text="Search Patient",font=("calibri",17,"bold")).place(x=400,y=150)
        txt_search=ctk.CTkEntry(self.root,textvariable=self.var_searchtxt,font=("goudy old style",15),fg_color="lightyellow").place(x=540,y=150)
        btn_search=ctk.CTkButton(self.root,command=self.search,text="Search",font=("goudy old style",18,"bold"),fg_color="#4caf50",width=150,height=30).place(x=700,y=150)
        
        #===============content=====================
        self.current_date=datetime.now().strftime("%d.%m.%Y")
        lbl_date=ctk.CTkLabel(self.root,text="Date",font=("goudy old style",17)).place(x=50,y=190)
        self.txt_date = PlaceholderEntry(self.root, placeholder_text=self.current_date, font=("goudy old style",15,"bold"), fg_color="lightyellow", width=180)
        self.txt_date.place(x=100, y=190)
        
        #============row1=============
        lbl_name=ctk.CTkLabel(self.root,text="Full Name",font=("goudy old style",17)).place(x=50,y=240)
        lbl_address=ctk.CTkLabel(self.root,text="Address",font=("goudy old style",17)).place(x=360,y=240)
        lbl_phone=ctk.CTkLabel(self.root,text="Phone No.",font=("goudy old style",17)).place(x=660,y=240)
             
        txt_name=ctk.CTkEntry(self.root,textvariable=self.var_name,font=("goudy old style",15),fg_color="lightyellow",width=180).place(x=150,y=240)
        txt_address=ctk.CTkEntry(self.root,textvariable=self.var_address,font=("goudy old style",15),fg_color="lightyellow",width=180).place(x=460,y=240)
        txt_phone=ctk.CTkEntry(self.root,textvariable=self.var_phone,font=("goudy old style",15),fg_color="lightyellow",width=180).place(x=760,y=240)
       
        #===================row2=================
         
        lbl_prof=ctk.CTkLabel(self.root,text="Profession",font=("goudy old style",17)).place(x=50,y=290)
        lbl_dob=ctk.CTkLabel(self.root,text="Age",font=("goudy old style",17)).place(x=360,y=290)
        lbl_gender=ctk.CTkLabel(self.root,text="Gender",font=("goudy old style",17)).place(x=660,y=290)
        
        txt_prof=ctk.CTkEntry(self.root,textvariable=self.var_prof,font=("goudy old style",15),fg_color="lightyellow",width=180).place(x=150,y=290)
        txt_dob=ctk.CTkEntry(self.root,textvariable=self.var_dob,font=("goudy old style",15),fg_color="lightyellow",width=180).place(x=410,y=290)
        self.cmb_gender = ctk.CTkComboBox(self.root, values=("Select", "Male", "Female"),fg_color="lightyellow", state='readonly', justify=CENTER, font=("calibri",15), width=180)
        self.cmb_gender.set("Select")
        self.cmb_gender.place(x=760, y=290)
        #===================row3=================
        
        lbl_mc=ctk.CTkLabel(self.root,text="Main Complain",font=("goudy old style",17)).place(x=50,y=340)
        lbl_doctor=ctk.CTkLabel(self.root,text="Doctor",font=("goudy old style",17)).place(x=360,y=340)
        
        txt_mc=ctk.CTkEntry(self.root,textvariable=self.var_mc,font=("goudy old style",15),fg_color="lightyellow",width=180).place(x=170,y=340)
        btn_doctor=ctk.CTkButton(self.root,text="Select",font=("goudy old style",12,"bold"),command=self.show_doctor_listbox,corner_radius=0,fg_color="#2ecc71", hover_color="#f1c40f",width=110)
        btn_doctor.place(x=450,y=340)
        self.lsb_doc = tk.Listbox(self.root, selectmode="multiple", exportselection=0)
        self.lsb_doc.place(x=-500, y=-500)  # Place the listbox outside the visible area
        self.lsb_doc.bind("<<ListboxSelect>>", self.update_doc_combobox)

        for value in self.doc_list:
            self.lsb_doc.insert(tk.END, value)
            
        lbl_tooth=ctk.CTkLabel(self.root,text="Tooth",font=("goudy old style",17)).place(x=660,y=340)
        self.tooth_type = ttk.Combobox(self.root,textvariable=self.var_tooth,values=["Select","Primary","Permanent"],state='readonly', justify=CENTER)
        self.tooth_type.set(["Select"])
        self.tooth_type.bind("<<ComboboxSelected>>", self.on_tooth_type_selected)
        self.tooth_type.place(x=760,y=340,width=200)

        self.tooth_selection = tk.Listbox(self.root,selectmode="multiple",exportselection=0, state='disabled')  # Initially disable the listbox
        self.tooth_selection.bind("<<ListboxSelect>>", self.on_tooth_selected)  # Bind the selection event
        self.tooth_selection.place(x=960,y=300,height=100)

                
        
        
        #===================row4=================
        
        lbl_obs=ctk.CTkLabel(self.root,text="Observations",font=("goudy old style",17)).place(x=50,y=380)
        lbl_tp=ctk.CTkLabel(self.root,text="Treatment History",font=("goudy old style",17)).place(x=500,y=380)
        
        self.txt_obs = ctk.CTkTextbox(self.root, font=("goudy old style", 15), fg_color="lightyellow", width=300, height=60,border_width=1)
        self.txt_obs.place(x=170,y=380)

        self.txt_tp = ctk.CTkTextbox(self.root, font=("goudy old style", 15), fg_color="lightyellow", width=300, height=60,border_width=1)
        self.txt_tp.place(x=640,y=380)
        
       
        
        #=============Buttons============
        btn_add=ctk.CTkButton(self.root,text="Save",command=self.add,font=("arial",15),fg_color="#2196f3",width=110,height=28).place(x=500,y=450)
        btn_update=ctk.CTkButton(self.root,text="Update",command=self.update,font=("arial",15),fg_color="#130f40",width=110,height=28).place(x=620,y=450)
        btn_delete=ctk.CTkButton(self.root,text="Delete",command=self.delete,font=("arial",15),fg_color="#f44336",width=110,height=28).place(x=740,y=450)
        btn_clear=ctk.CTkButton(self.root,text="Clear",command=self.clear,font=("arial",15),fg_color="#607d8b",width=110,height=28).place(x=860,y=450)
        btn_save=ctk.CTkButton(self.root,text="Save Rec",command=self.other,font=("arial",15),fg_color="#6ab04c",width=110,height=28).place(x=980,y=450)
        btn_view=ctk.CTkButton(self.root,text="View Record",command=self.view_rec,font=("arial",15),fg_color="#3c40c6",width=130,height=28).place(x=1100,y=450)
        btn_print=ctk.CTkButton(self.root,text="Print Record",command=self.print_rec,font=("arial",15),fg_color="#ff5e57",width=110,height=28).place(x=1240,y=450)
       
        #=====================Employee Details============
        pat_frame=ctk.CTkFrame(self.root)
        pat_frame.place(x=10,y=490)
        
        scolly=Scrollbar(pat_frame,orient=VERTICAL)
        scollx=Scrollbar(pat_frame,orient=HORIZONTAL)
        
        self.PatientTable=ttk.Treeview(pat_frame,columns=("pat_id","name","doctor_name","address","phone","profession","dob","gender","mc","tooth","observations","tp","date"),yscrollcommand=scolly.set,xscrollcommand=scollx.set)
        scollx.pack(side=BOTTOM,fill=X)
        scolly.pack(side=RIGHT,fill=Y)
        scollx.config(command=self.PatientTable.xview)
        scolly.config(command=self.PatientTable.yview)
        
        self.PatientTable.heading("pat_id",text="PAT ID")
        self.PatientTable.heading("name",text="Full Name")
        self.PatientTable.heading("doctor_name",text="Doctor")
        self.PatientTable.heading("address",text="Address")
        self.PatientTable.heading("phone",text="Phone")
        self.PatientTable.heading("profession",text="Profession")
        self.PatientTable.heading("dob",text="Age")
        self.PatientTable.heading("gender",text="gender")
        self.PatientTable.heading("mc",text="Main Complain")
        self.PatientTable.heading("tooth",text="Teeth")
        self.PatientTable.heading("observations",text="Observations")
        self.PatientTable.heading("tp",text="Treatment History")
        self.PatientTable.heading("date",text="Date")
        self.PatientTable["show"] ="headings" 
        self.PatientTable.column("pat_id",width=50)
        self.PatientTable.column("name",width=100)
        self.PatientTable.column("doctor_name",width=100)
        self.PatientTable.column("address",width=100)
        self.PatientTable.column("phone",width=130)
        self.PatientTable.column("profession",width=100)
        self.PatientTable.column("dob",width=100)
        self.PatientTable.column("gender",width=60)
        self.PatientTable.column("mc",width=100)
        self.PatientTable.column("tooth",width=100)
        self.PatientTable.column("observations",width=150)
        self.PatientTable.column("tp",width=150)
        self.PatientTable.column("date",width=100)
        self.PatientTable.pack(fill=BOTH,expand=1)
        self.PatientTable.bind("<ButtonRelease-1>",self.get_data)
        permission.interact_with_database((resource_path('PRMS.db'))) 
        self.show()
        
        
        
        
        #==================Footer================
        date_=time.strftime("%Y")
        lbl_footer=ctk.CTkLabel(self.root,text=f"Copyright @ {date_} RootTech", font=("times new roman",12,"bold")).pack(side=BOTTOM,fill=X)
        
        permission.interact_with_database((os.path.join(os.getcwd(),resource_path('PRMS.db'))) )
        self.txt_date.configure(placeholder_text=self.current_date)
        self.update_date_time()
        #=====================All functions=================
   
            
    def on_tooth_type_selected(self, event):
        self.tooth_selection['state'] = 'normal'  # Enable the listbox when a selection is made
        self.tooth_selection.delete(0, tk.END)
        tooth_type = self.tooth_type.get()
        if tooth_type == "Primary":
            self.teeth = ["#51", "#52", "#53","#54","#55","#61","#61","#62","#63",
                        "#64","#65","#71","#72","#73","#74","#75","#81","#82","#83","#84","#85"]  
        else:
            self.teeth = ["#18","#17","#16","#15","#14","#13","#12","#11",
                        "#21","#22","#23","#24","#25","#26","#27","#28","#48",
                        "#47","#46","#45","#44","#43","#42","#41","#31","#32","#33",
                        "#34","#35","#36","#37","#38"] 
        for tooth in self.teeth:
            self.tooth_selection.insert(tk.END, tooth)

    def on_tooth_selected(self, event):  # New function to handle selection event
        selected_teeth = [self.tooth_selection.get(idx) for idx in self.tooth_selection.curselection()]
        self.var_tooth.set(", ".join(selected_teeth))
        
   
    
    def show_doctor_listbox(self):
        screen_width=self.root.winfo_screenwidth()
        screen_height=self.root.winfo_screenheight()
        if self.lsb_doc.winfo_x() < 0:  # If the listbox is outside the visible area
            self.lsb_doc.place(x=screen_width-760,y=screen_height-420, width=200)  # Move it to the desired position
            self.lsb_doc.lift()  # Bring the widget to the top of the stacking order
        else:
            self.lsb_doc.place(x=-500, y=-500)  # Move it outside the visible area
     
                
    def update_doc_combobox(self, event=None):
        selected_values = [self.lsb_doc.get(idx) for idx in self.lsb_doc.curselection()]
        self.var_doctor.set(", ".join(selected_values))
   
        
    def fetch_doctor(self):
        permission.interact_with_database((resource_path('PRMS.db'))) 
        self.doc_list.append("Empty")
        con = sqlite3.connect(database=os.path.join(os.getcwd(),resource_path(r'PRMS.db')))
        cur = con.cursor()
        try:
            cur.execute("Select doc_name from doctor")
            cat = cur.fetchall()
            if cat:
                self.doc_list = ["Select"] + [i[0] for i in cat]
        except Exception as ex:
            messagebox.showerror("Error", f"Error due to : {str(ex)}", parent=self.root)
            
            
    def add(self):
        permission.interact_with_database((resource_path('PRMS.db'))) 
        con=sqlite3.connect(database=os.path.join(os.getcwd(),resource_path(r'PRMS.db')))
        cur=con.cursor()
        try:
            if self.var_name.get()=="":
                messagebox.showerror("Error","Patient Name is required",parent=self.root)
            else:
                cur.execute("Select * from archives where name=?",(self.var_name.get(),))
                row=cur.fetchone()
                if row!=None:
                    messagebox.showerror("Error","This Patient has already been Archived, try a different",parent=self.root)
                else:
                    cur.execute("Insert into archives (name,doctor_name,address,phone,profession,dob,gender,mc,tooth,observations,tp,date) values(?,?,?,?,?,?,?,?,?,?,?,?)",(
                                                self.var_name.get(),
                                                self.var_doctor.get(),                                                
                                                self.var_address.get(),
                                                self.var_phone.get(),
                                                self.var_prof.get(),                                                
                                                self.var_dob.get(),
                                                self.cmb_gender.get(),                                                
                                                self.var_mc.get(),
                                                self.var_tooth.get(),
                                                self.txt_obs.get('1.0',END),
                                                self.txt_tp.get('1.0',END), 
                                                self.txt_date.get()              
                        
                    ))
                    con.commit()
                    messagebox.showinfo("Success","Patient Record Added Successfully",parent=self.root)
                    self.show()
                    self.clear()
        except Exception as ex:
            messagebox.showerror("Error",f"Error due to : {str(ex)}",parent=self.root)
            
    def show(self):
        permission.interact_with_database((resource_path('PRMS.db'))) 
        con=sqlite3.connect(database=os.path.join(os.getcwd(),resource_path(r'PRMS.db')))
        cur=con.cursor()
        try:
            cur.execute("select * from archives")
            rows=cur.fetchall()
            self.PatientTable.delete(*self.PatientTable.get_children())
            for row in rows:
                self.PatientTable.insert('',END,values=row)
        except Exception as ex:
            messagebox.showerror("Error",f"Error due to : {str(ex)}",parent=self.root)
    
    def get_data(self,ev):
        f=self.PatientTable.focus()
        content=(self.PatientTable.item(f))  
        row=content['values'] 
        self.var_pat_id.set(row[0])
        self.var_name.set(row[1])  
        self.var_doctor.set(row[2])                                            
        self.var_address.set(row[3])
        self.var_phone.set(row[4])
        self.var_prof.set(row[5])                                                
        self.var_dob.set(row[6])
        self.cmb_gender.set(row[7])                                                
        self.var_mc.set(row[8])
        self.var_tooth.set(row[9])
        self.txt_obs.delete('1.0',END)
        self.txt_obs.insert(END,row[10])
        self.txt_tp.delete('1.0',END)
        self.txt_tp.insert(END,row[11])
        self.txt_date.configure(placeholder_text=f"{row[12]}")
        
        
        
    def update(self):
        if not self.var_name.get():
            messagebox.showerror("Error", "Patient Name is required", parent=self.root)
            return
        permission.interact_with_database((resource_path('PRMS.db')))
        with sqlite3.connect(database=resource_path('PRMS.db')) as con:
            cur = con.cursor()
            params = (self.var_pat_id.get(),)
            cur.execute("SELECT * FROM archives WHERE pat_id=?", params)
            
            if cur.fetchone() is None:
                messagebox.showerror("Error", "No patient record found with the given name", parent=self.root)
                return

            params = (
                self.var_name.get(),
                self.var_doctor.get(),
                self.var_address.get(),
                self.var_phone.get(),
                self.var_prof.get(),
                self.var_dob.get(),
                self.cmb_gender.get(),
                self.var_mc.get(),
                self.var_tooth.get(),
                self.txt_obs.get('1.0', END),
                self.txt_tp.get('1.0',END), 
                self.txt_date.get(),
                self.var_pat_id.get()
            )
            cur.execute("UPDATE archives SET name=?, doctor_name=?, address=?, phone=?, profession=?, dob=?, gender=?, mc=?, tooth=?, observations=?, tp=?, date=? WHERE pat_id=?", params)
            con.commit()
            messagebox.showinfo("Success", "Patient Archived Record Updated Successfully", parent=self.root)
            self.show()
 
         
      
    def delete(self):
        permission.interact_with_database((resource_path('PRMS.db'))) 
        con=sqlite3.connect(database=os.path.join(os.getcwd(),resource_path(r'PRMS.db')))
        cur=con.cursor()
        try:
            if self.var_name.get()=="":
                messagebox.showerror("Error","Patient Name is required",parent=self.root)
            else:
                cur.execute("Select * from archives where name=?",(self.var_name.get(),))
                row=cur.fetchone()
                if row==None:
                    messagebox.showerror("Error","Invalid Patient Name",parent=self.root)
                else:
                    op=messagebox.askyesno("Confirm",f"Do you really want to delete {self.var_name.get()} Archived Record ?",parent=self.root)
                    if op==True:                        
                        cur.execute("delete from archives where name=?",(self.var_name.get(),))
                        con.commit()
                        messagebox.showinfo("Delete","Patient Record Deleted Successfully",parent=self.root)
                        self.clear()                
        except Exception as ex:
            messagebox.showerror("Error",f"Error due to : {str(ex)}",parent=self.root)
     
    def clear(self):
        self.var_name.set("") 
        self.var_doctor.set("")                                             
        self.var_address.set("")
        self.var_phone.set("")
        self.var_prof.set("")                                                
        self.var_dob.set("")
        self.cmb_gender.set("Select") 
        self.txt_obs.delete('1.0',END)
        self.txt_tp.delete('1.0',END)
        self.var_mc.set("")
        self.var_tooth.set("Select")
        self.var_searchtxt.set("")
        self.txt_date.configure(placeholder_text=self.current_date)
        self.show()    
        
    def search(self):
        permission.interact_with_database((resource_path('PRMS.db'))) 
        con=sqlite3.connect(database=os.path.join(os.getcwd(),resource_path(r'PRMS.db')))
        cur=con.cursor()
        try:
            if self.var_searchtxt.get()=="":
                messagebox.showerror("Error","Search input is required",parent=self.root)
            else:        
                cur.execute("select pat_id,name,doctor_name,address,phone,profession,dob,gender,mc,tooth,observations,tp from archives where name LIKE '%"+self.var_searchtxt.get()+"%'")  
                rows=cur.fetchall()
                if len(rows)!=0:
                    self.PatientTable.delete(*self.PatientTable.get_children())
                    for row in rows:
                        self.PatientTable.insert('',END,values=row)
                else:
                    messagebox.showerror("Error","No record found!!!",parent=self.root)
        except Exception as ex:
            messagebox.showerror("Error",f"Error due to : {str(ex)}",parent=self.root)   
    
    def find_or_create_table(self,doc):
                for element in doc.element.body:
                    if isinstance(element, docx.oxml.table.CT_Tbl):
                        return docx.table.Table(element, doc)  # Use the existing table

                    # Create a new table with headers
                table = doc.add_table(rows=1, cols=6)
                table.style = 'Table Grid'
                table.autofit = True
                table.allow_autofit = True
                header_row = table.rows[0]
                header_row.cells[0].text = 'DATE'
                header_row.cells[1].text = 'TOOTH'
                header_row.cells[2].text = 'INTERVENTION HISTORY'
                header_row.cells[3].text = 'Amount DUE'
                header_row.cells[4].text = 'Amount PAID'
                header_row.cells[5].text = 'BALANCE'

                return table
        
    def generate_rec(self):
        self.existing_doc_path=os.path.join(os.getcwd(),resource_path(f'archives\\{self.var_name.get()}.docx'))
        if os.path.exists(self.existing_doc_path):
            doc=docx.Document(self.existing_doc_path) 
            table = self.find_or_create_table(doc)
            # Add a new row with data
            new_row = table.add_row().cells
            new_row[0].text = self.txt_date.get()
            new_row[1].text = self.var_tooth.get()  # Tooth
            new_row[2].text = self.txt_tp.get('1.0',END)  # Nature of intervention
            new_row[3].text = self.amt_due.get()  # Due
            new_row[4].text = self.amt_rcd.get()  # Paid
            new_row[5].text = self.amt_ble.get()  # Balance
            
            doc.save(os.path.join(os.getcwd(),resource_path(self.existing_doc_path)))
            

            # Show success message
            messagebox.showinfo("Success", "Patient Record has been Updated and Generated. Open the Record folder to view and print.", parent=self.root)
            
            return
        else:
            doc=docx.Document()
        header=doc.sections[0].header
        for paragraph in header.paragraphs:
            for run in paragraph.runs:
                run.clear()
        htable=header.add_table(1,2,Inches(6))
        htab_cells=htable.rows[0].cells
        
        ht0=htab_cells[0].add_paragraph()
        logo=ht0.add_run()
        logo.add_picture(resource_path('images\\emma.png'))
        ht1=htab_cells[1].add_paragraph(f'''\tEMMANUEL DENTAL CLINIC 
        BONAMOUSSADI
        Pr.(Dr) AGBOR MICHEAL ASHU
        Tel: 677 17 01 67/697 12 27 82
        www.emmanueldentalcare.org
        ''')
        ht1.alignment=WD_ALIGN_PARAGRAPH.CENTER
        
        title=doc.add_heading(f'''\tPATIENT RECORD
''',level=0)
        title.alignment=WD_ALIGN_PARAGRAPH.CENTER
        
        #patient Details head
        doc.add_heading('PATIENT DETAILS')
        #patient Details
        details=doc.add_paragraph()
        details.add_run('Full Name: ')
        details.add_run(f"{self.var_name.get()}")
        details.add_run('\n')
        details.add_run('Address: ')
        details.add_run(f"{self.var_address.get()}")
        details.add_run('\n')
        details.add_run('Phone: ')
        details.add_run(f"{self.var_phone.get()}")
        details.add_run('\n')
        details.add_run('Age: ')
        details.add_run(f"{self.var_dob.get()}")
        details.add_run('\n')
        
        #observations
        
        doc.add_heading('OBSERVATIONS')
        ob=doc.add_paragraph()
        ob.add_run('Doctor: ').bold=True
        ob.add_run(f'Dr.{self.var_doctor.get()}')
        ob.add_run('\n')
        ob.add_run('Main Complain: ').bold=True
        ob.add_run(f'{self.var_mc.get()}')
        ob.add_run('\n')
        ob.add_run('Treatment Plan: ').bold=True
        ob.add_run(f'{self.var_tp.get()}')
        ob.add_run('\n')
        ob.add_run('Observations: ').bold=True
        ob.add_run(f'{self.txt_obs.get('1.0',END)}')
        
        #table
        
        table = self.find_or_create_table(doc)

        # Add a new row with data
        new_row = table.add_row().cells
        new_row[0].text = self.txt_date.get()
        new_row[1].text = self.var_tooth.get() # Tooth
        new_row[2].text = self.txt_tp.get('1.0',END)  # History of intervention
        new_row[3].text = self.amt_due.get()  # Due
        new_row[4].text = self.amt_rcd.get()  # Paid
        new_row[5].text = self.amt_ble.get()  # Balance
        
        doc.save(os.path.join(os.getcwd(),resource_path(self.existing_doc_path)))
        # Show success message
        messagebox.showinfo("Success", "Patient Record has been Generated and|or Updated. Click View Rec button to View Patient Record", parent=self.root)
        
            
    def other(self):
        file_path=os.path.join(os.getcwd(),resource_path(f'archives\\{self.var_name.get()}.docx'))
        if not os.path.exists(file_path):
            if self.var_name.get()=='' or self.var_address.get()=='' or self.var_phone.get()=='':
                messagebox.showerror("Error","All Patient Details are required, Select patient Data from database bellow",parent=self.root)
            else:        
                self.gen_win=Toplevel(self.root)     
                self.gen_win.title('SmileScribePro. GENERATE PATIENT RECORD')
                self.gen_win.iconbitmap(resource_path('icon.ico')) 
                self.gen_win.geometry('450x380+500+100')     
                self.gen_win.focus_force()
                
                self.amt_due=StringVar()
                self.amt_rcd=StringVar()
                self.amt_ble=StringVar()
            
                
                title=Label(self.gen_win,text="ADD BILL DATA",font=('goudy old style',15,'bold'),bg="#3f51b5",fg="white").pack(side=TOP,fill=X)                   
                lbl_amt_due=Label(self.gen_win,text="AMOUNT DUE",font=("times new roman",15)).place(x=20,y=60)
                txt_amt_due=Entry(self.gen_win,textvariable=self.amt_due,font=("times new roman",15),bg='lightyellow').place(x=20,y=100,width=250,height=30)
                
                lbl_amt_rcd=Label(self.gen_win,text="AMOUNT RECEIVED",font=("times new roman",15)).place(x=20,y=160)
                txt_amt_rcd=Entry(self.gen_win,textvariable=self.amt_rcd,font=("times new roman",15),bg='lightyellow').place(x=20,y=190,width=250,height=30)
                
                
                lbl_amt_ble=Label(self.gen_win,text="AMOUNT LEFT",font=("times new roman",15)).place(x=20,y=225)
                txt_amt_ble=Entry(self.gen_win,textvariable=self.amt_ble,font=("times new roman",15),bg='lightyellow').place(x=20,y=260,width=250,height=30)
                
                lbl_note=Label(self.gen_win,text=f"\t\tNote:'Enter 0 in AMOUNT DUE \n\tIF PATIENT HAS PAID ALL HIS/HER BILL'",font=("goudy old style",12),anchor='w',bg="white",fg="red").pack(side=BOTTOM,fill=X)    
                
                self.btn_update=Button(self.gen_win,text="SAVE | UPDATE RECORD",command=self.add_doc_pat_rec,font=("times new roman",15),bg='lightblue')
                self.btn_update.place(x=100,y=300,width=250,height=30)
        else:
            op=messagebox.askyesno('confirm',f"{file_path} already exist, Do you want to update Archived file?",parent=self.root)
            if op==True:
                self.gen_win=Toplevel(self.root)     
                self.gen_win.title('SmileScribePro. GENERATE PATIENT ARCHIVE(DOCUMENT)')
                self.gen_win.iconbitmap(resource_path('icon.ico')) 
                self.gen_win.geometry('450x380+500+100')     
                self.gen_win.focus_force()
                
                self.amt_due=StringVar()
                self.amt_rcd=StringVar()
                self.amt_ble=StringVar()
            
                
                title=Label(self.gen_win,text="ADD BILL DATA",font=('goudy old style',15,'bold'),bg="#3f51b5",fg="white").pack(side=TOP,fill=X)                   
                lbl_amt_due=Label(self.gen_win,text="AMOUNT DUE",font=("times new roman",15)).place(x=20,y=60)
                txt_amt_due=Entry(self.gen_win,textvariable=self.amt_due,font=("times new roman",15),bg='lightyellow').place(x=20,y=100,width=250,height=30)
                
                lbl_amt_rcd=Label(self.gen_win,text="AMOUNT RECEIVED",font=("times new roman",15)).place(x=20,y=160)
                txt_amt_rcd=Entry(self.gen_win,textvariable=self.amt_rcd,font=("times new roman",15),bg='lightyellow').place(x=20,y=190,width=250,height=30)
                
                
                lbl_amt_ble=Label(self.gen_win,text="AMOUNT LEFT",font=("times new roman",15)).place(x=20,y=225)
                txt_amt_ble=Entry(self.gen_win,textvariable=self.amt_ble,font=("times new roman",15),bg='lightyellow').place(x=20,y=260,width=250,height=30)
                
                lbl_note=Label(self.gen_win,text=f"\t\tNote:'Enter 0 in AMOUNT DUE \n\tIF PATIENT HAS PAID ALL HIS/HER BILL'",font=("goudy old style",12),anchor='w',bg="white",fg="red").pack(side=BOTTOM,fill=X)    
                
                self.btn_update=Button(self.gen_win,text="SAVE | UPDATE RECORD",command=self.add_doc_pat_rec,font=("times new roman",15),bg='lightblue')
                self.btn_update.place(x=100,y=300,width=250,height=30)
            
    
    def add_doctor_patient_rec(self):
        if not self.var_doctor.get():
            messagebox.showerror("Error", "Doctor Name is required", parent=self.root)
            return
        try:
            with sqlite3.connect(database=os.path.join(os.getcwd(),resource_path(r'PRMS.db'))) as con:
                cur = con.cursor()
                # Check if a record with the same doctor and patient name already exists
                cur.execute("SELECT doc_id, intervention, amount_paid, date FROM doctor_patient_records WHERE doc_name=? AND pat_name=?", (self.var_doctor.get(), self.var_name.get()))
                existing_record = cur.fetchone()

                # Fetch tp and date from the patient table
                cur.execute("SELECT name, doctor_name, tp, date FROM patient WHERE name=?", (self.var_name.get(),))
                patient_record = cur.fetchone()
                if patient_record:
                    name, doctor_name, tp, date = patient_record

                    # Fetch doc_id from the doctor table
                    cur.execute("SELECT doc_id FROM doctor WHERE doc_name=?", (doctor_name,))
                    doc_id_row = cur.fetchone()
                    doc_id = doc_id_row[0] if doc_id_row else None

                    # If the record exists and the data is different, update the record
                    if existing_record and (existing_record[1] != tp or existing_record[2] != self.amt_rcd.get() or existing_record[3] != date):
                        params = (doc_id, doctor_name, name, tp, self.amt_rcd.get(), date, self.var_doctor.get(), self.var_name.get())
                        cur.execute("UPDATE doctor_patient_records SET doc_id=?, doc_name=?, pat_name=?, intervention=?, amount_paid=?, date=? WHERE doc_name=? AND pat_name=?", params)
                        messagebox.showinfo("Success", "Record updated successfully", parent=self.root)
                    # If the record does not exist, insert a new record
                    elif not existing_record:
                        params = (doc_id, doctor_name, name, tp, self.amt_rcd.get(), date)
                        cur.execute("INSERT INTO doctor_patient_records (doc_id, doc_name, pat_name, intervention, amount_paid, date) VALUES (?, ?, ?, ?, ?, ?)", params)
                        messagebox.showinfo("Success", "New record added successfully", parent=self.root)
                else:
                    messagebox.showerror("Error", "No patient record found with the given name", parent=self.root)
                con.commit()
        except Exception as ex:
            messagebox.showerror("Error",f"Error due to : {str(ex)}",parent=self.root)   
    

    
    
    '''def add_doctor_patient_rec(self):
        if not self.var_doctor.get():
            messagebox.showerror("Error", "Doctor Name is required", parent=self.root)
            return

        with sqlite3.connect(database=resource_path(r'PRMS.db')) as con:
            cur = con.cursor()
            params = (self.var_doctor.get(), self.var_name.get())
            cur.execute("SELECT * FROM doctor_patient_records WHERE doc_name=? AND pat_name=?", params)
            
            row = cur.fetchone()
            if row is not None:
                cur.execute("SELECT name, doctor_name, tp, date FROM patient WHERE name=?", (self.var_name.get(),))
                row = cur.fetchone()
                if row is not None:
                    name, doctor_name, tp, date = row
                    existing_data = (row[3], row[4], row[5])  # Assuming these are intervention, amount_paid, date in your table
                    new_data = (tp, self.amt_rcd.get(), date)
                    if existing_data != new_data:
                        cur.execute("SELECT doc_id FROM doctor WHERE doc_name=?", (self.var_doctor.get(),))
                        row = cur.fetchone()
                        doc_id = row[0] if row else None
                        params = (doc_id, doctor_name, name, tp, self.amt_rcd.get(), date)
                        cur.execute("UPDATE doctor_patient_records SET doc_id=?, doc_name=?, pat_name=?, intervention=?, amount_paid=?, date=? WHERE doc_name=? AND pat_name=?", params)
                        messagebox.showinfo("Success", "Record updated successfully", parent=self.root)
                    else:
                        messagebox.showinfo("No Update Required", "The existing record has the same data.", parent=self.root)
            else:
                cur.execute("SELECT name, doctor_name, tp, date FROM patient WHERE pat_id=?", (self.var_pat_id.get(),))
                row = cur.fetchone()
                if row is not None:
                    name, doctor_name, tp, date = row
                    cur.execute("SELECT doc_id FROM doctor WHERE doc_name=?", (doctor_name,))
                    doc_id_row = cur.fetchone()
                    doc_id = doc_id_row[0] if doc_id_row else None
                    params = (doc_id, doctor_name, name, tp, self.amt_rcd.get(), date)
                    cur.execute("INSERT INTO doctor_patient_records (doc_id, doc_name, pat_name, intervention, amount_paid, date) VALUES (?, ?, ?, ?, ?, ?)", params)
                    messagebox.showinfo("Success", "New record added successfully", parent=self.root)
                    con.commit()'''


                
    def add_doc_pat_rec(self): 
            if self.amt_due.get()=='0':
                self.generate_rec() 
                return self.gen_win.destroy() 
            elif self.amt_due.get()=='' or self.amt_ble.get()=='' or self.amt_rcd.get()=='':
                messagebox.showerror("Error","Complete Patient Bill Info",parent=self.gen_win)
            else:
                self.generate_rec()
                self.gen_win.destroy() 
  
                
            
            
    def view_rec(self):
        file_path=os.path.join(os.getcwd(),resource_path(f'archives\\{self.var_name.get()}.docx'))
        if not os.path.exists(file_path):
            messagebox.showerror("View Error",f"The file {file_path} does not exist.\nSave Rec then attempt Viewing",parent=self.root)
            
        else:
            notification.notify(title="Success",message="Opening file...!!! be patient...",timeout=30,parent=self.root)
            # Open the file in default program
            os.startfile(file_path)
     
            
            
        
        '''self.view_win=Toplevel(self.root)     
        self.view_win.title('GENERATE PATIENT RECORD')
        self.view_win.iconbitmap(resource_path('icon.ico')) 
        self.view_win.geometry('1100x500+220+130')     
        self.view_win.focus_force()      
        
        V_RTitle=Label(self.view_win,text="View Patient Consultation Sheet Area",font=("goudy old style",20,"bold"),bg="#f44336",fg="white").pack(side=TOP,fill=X)
        scrolly=Scrollbar(self.view_win,orient=VERTICAL)
        scrollx=Scrollbar(self.view_win,orient=HORIZONTAL)
        scrolly.pack(side=RIGHT,fill=Y)
        scrollx.pack(side=BOTTOM,fill=X)
        
        self.txt_file_area=Text(self.view_win,font=("Courier",18,"bold"),wrap=WORD,yscrollcommand=scrolly.set,xscrollcommand=scrollx.set)
        self.txt_file_area.pack(fill=BOTH,expand=1)
        scrolly.config(command=self.txt_file_area.yview)
        scrollx.config(command=self.txt_file_area.xview) 
         
        self.docx_text=self.read_docx(self.existing_doc_path)
        self.txt_file_area.insert('1.0',self.docx_text)'''    
       
        
    def read_docx(self,file_path):
        doc=docx.Document(resource_path(file_path))
        fulltext=[]
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                fulltext.append(paragraph.text)
        for para in doc.paragraphs:
            fulltext.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    fulltext.append(cell.text)
        return '\n'.join(fulltext)      
    
   
    def print_rec(self):
        file_path=os.path.join(os.getcwd(),resource_path(f'archives\\{self.var_name.get()}.docx'))
        if not os.path.exists(file_path):
            messagebox.showerror("Print Error",f"The file {file_path} does not exist.\nSave Rec then attempt printing",parent=self.root)
        else:
            notification.notify(title="Success",message="Printing...!!!",timeout=30,parent=self.root)
            doc=docx.Document(file_path)
            printer_name=win32print.GetDefaultPrinter()
            win32api.ShellExecute(0,"print",file_path,'d:"%s"'%printer_name,".",0)
    
        
    
    def update_date_time(self):   
        time_=time.strftime("%H:%M:%S")
        date_=time.strftime("%d:%m:%Y")
        self.lbl_clock.configure(text=f"Welcome to SmileScribe Professional Patients Record Management System\t\t Date: {str(date_)}\t\t Time: {str(time_)}")
        self.lbl_clock.after(200,self.update_date_time)   
        
        
    def exit(self):
        self.is_running = False
        self.root.destroy() 
    
            
    
        
        
if __name__=="__main__":
    root=ctk.CTk()
    obj=archiveClass(root)
    root.mainloop()