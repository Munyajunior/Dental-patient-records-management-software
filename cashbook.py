from tkinter import*
import customtkinter as ctk
from PIL import Image
from tkinter import ttk,messagebox
import sqlite3
import time
import docx
from docx import Document
from plyer import notification
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import win32print
import win32api
import docx
from num2words import num2words
from docx import Document
from docx.enum.text import  WD_ALIGN_PARAGRAPH
from datetime import datetime
import os
import permission
import sys


def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS2
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)



ctk.set_appearance_mode("system")
ctk.set_default_color_theme("blue")


class cashBookClass(ctk.CTk):
    def __init__(self,root):
        self.root=root
        self.root.state('zoomed')
        # Bind the exit method to the window's close button
        self.root.protocol("WM_DELETE_WINDOW", self.exit)
        screen_width=self.root.winfo_screenwidth()
        screen_height=self.root.winfo_screenheight()
        self.root.geometry(f'{screen_width}x{screen_height}')
        self.root.iconbitmap(resource_path('icon.ico'))
        self.root.title("SmileScribePro")
        self.quote_list=[]
        self.quote_data_list=[]
        permission.interact_with_database((resource_path('PRMS.db')))
        self.con = sqlite3.connect(database=resource_path(r'PRMS.db'))
        self.cur = self.con.cursor()
               
        #============================Title======================
        self.icon_title=ctk.CTkImage(dark_image=Image.open(resource_path("images\\logo.png")),
                                    light_image=Image.open(resource_path("images\\logo.png")),size=(50,50))
        title = ctk.CTkLabel(
        self.root,
        text="SmileScribePro",
        image=self.icon_title,
        compound=LEFT,
        font=("sans-serif", 40, "bold"),
        fg_color=("#3498db","#fff"),
        text_color=("#fff","#3498db"),
        anchor='w',
        width=screen_width,  # Set the width to be the screen width
        height=68,padx=20
    ).place(anchor='nw', x=0, y=0)
        #=============Btn_Logout================================
        # Calculate the position of the button based on the screen size
        x_position = screen_width - 100  # Adjust this value as needed
        y_position = screen_height - 735  # Adjust this value as needed
        btn_logout = ctk.CTkButton(self.root,text="Exit"
                                    ,command=self.exit
                                    ,font=("calibri",20,"bold")
                                    ,fg_color=("#273c75","#353b48")
                                    ,hover_color="#e84118"
                                    ,height=50,width=150
                                    ,corner_radius=0
                                    ).place(anchor='e', x=x_position, y=30)
        #======================Clock============================
        self.lbl_clock=ctk.CTkLabel(self.root,text="Welcome to SmileScribePro - Professional Patient Records Management System\t\t Date: DD-MM-YYYY\t\t Time: HH:MM:SS"
                                    ,font=("calibri",18,"bold")
                                    ,fg_color=("#487eb0","#4d636d"),bg_color="#fff"
                                    ,width=screen_width
                                    ,text_color="#fff")
        self.lbl_clock.place(x=0,y=65)

        #=======================Content=======================
        #===================title=========
        title=ctk.CTkLabel(self.root,text="CASH BOOK",font=("goudy old style",20,"bold"),fg_color="#0f4d7d",text_color="#fff",width=1000).place(x=170,y=100)
       
       #=================Doctor Records Main frame================
        DoctorFrame = ctk.CTkFrame(self.root)
        DoctorFrame.place(x=170,y=130) 

        pTitle=ctk.CTkLabel(DoctorFrame,text="All Patient Records",font=("arial",20,"bold"),text_color="#fff",fg_color="#262626",width=380).pack(side=TOP,fill=X)

        #=================Doctor Record Search frame================
        
        #==========Variables================
        self.var_search=ctk.StringVar()

        lbl_search=ctk.CTkLabel(self.root,text="Search Patient Records | By Name ",font=("times new roman",15,"bold")).place(x=170,y=170)

        lbl_search=ctk.CTkLabel(self.root,text="Records",font=("times new roman",15,"bold")).place(x=170,y=200)
        txt_search=ctk.CTkEntry(self.root,textvariable=self.var_search,font=("times new roman",15),fg_color="lightyellow",width=150,height=22).place(x=260,y=200)
        btn_search=ctk.CTkButton(self.root,text="Search",command=self.search,font=("arial",15),fg_color="#2196f3",width=100,height=25).place(x=460,y=200)
        btn_show_all=ctk.CTkButton(self.root,text="Show All",command=self.show,font=("arial",15),fg_color="#083531",width=100,height=25).place(x=460,y=170)

        #=====================Doctor Records Frame============
        CashBookFrame=ctk.CTkFrame(self.root)
        CashBookFrame.place(x=170,y=240)

        scolly=Scrollbar(CashBookFrame,orient=VERTICAL)
        scollx=Scrollbar(CashBookFrame,orient=HORIZONTAL)
        
        def treeview_sort_column(tv, col, reverse):            
            if col == 'DATE':
                l = [(datetime.strptime(tv.set(k, col), "%d.%m.%Y"), k) for k in tv.get_children('')]
            else:
                l = [(tv.set(k, col), k) for k in tv.get_children('')]
            l.sort(reverse=reverse)
            for index, (val, k) in enumerate(l):
                tv.move(k, '', index)
            tv.heading(col, command=lambda _col=col: treeview_sort_column(tv, _col, not reverse))

        self.Cash_BookTable=ttk.Treeview(CashBookFrame,columns=("pat_name","intervention", "amount_paid", "date"),yscrollcommand=scolly.set,xscrollcommand=scollx.set)
        scolly.pack(side=RIGHT,fill=Y)
        scollx.pack(side=BOTTOM,fill=X)
        scolly.config(command=self.Cash_BookTable.yview)
        scollx.config(command=self.Cash_BookTable.xview)

        columns = {"pat_name":"PATIENT_NAME","intervention":"INTERVENTION", "amount_paid":"AMOUNT_PAID", "date":"DATE"}
        
        for k,v in columns.items():
            self.Cash_BookTable.heading(k, text=v, command=lambda _col=k: treeview_sort_column(self.Cash_BookTable, _col, False))

        self.Cash_BookTable["show"] ="headings" 
        self.Cash_BookTable.column("pat_name",width=100)
        self.Cash_BookTable.column("intervention",width=100)  
        self.Cash_BookTable.column("amount_paid",width=100)
        self.Cash_BookTable.column("date",width=100)          
        self.Cash_BookTable.pack(fill=BOTH,expand=1)
        self.Cash_BookTable.bind("<ButtonRelease-1>")
        
         #=======================Generate cashbook Widgets Frame===================
        #=====Variable=====
        self.month_=StringVar()
        self.year_=StringVar()
        
        
        Gen_CashBookWidgetFrame=ctk.CTkFrame(self.root,width=400,height=150)
        Gen_CashBookWidgetFrame.place(x=600,y=240)
        
        lbl_p_month=ctk.CTkLabel(Gen_CashBookWidgetFrame,text="Month",font=("times new roman",15)).place(x=5,y=5)       
        txt_p_month=ctk.CTkEntry(Gen_CashBookWidgetFrame,textvariable=self.month_,font=("times new roman",15)).place(x=5,y=35)
        
        lbl_p_yr=ctk.CTkLabel(Gen_CashBookWidgetFrame,text="Year",font=("times new roman",15)).place(x=5,y=70)       
        txt_p_yr=ctk.CTkEntry(Gen_CashBookWidgetFrame,textvariable=self.year_,font=("times new roman",15)).place(x=5,y=100)
        
        btn_sort_cashb=ctk.CTkButton(Gen_CashBookWidgetFrame,text="Sort CashBook",command=self.sort_treeTable,font=("times new roman",15,"bold"),fg_color="orange").place(x=200,y=15)      
        btn_gen_cashb=ctk.CTkButton(Gen_CashBookWidgetFrame,text="Generate CashBook",command=self.generate_cash_book,font=("times new roman",15,"bold"),fg_color="#009688").place(x=200,y=50)
        btn_print_cashb=ctk.CTkButton(Gen_CashBookWidgetFrame,text="print CashBook",command=self.print_cashbook,font=("times new roman",15,"bold"),fg_color="#009623").place(x=200,y=85)
        
        
        
        
        
        #==================Footer================
        date_=time.strftime("%Y")
        lbl_footer=ctk.CTkLabel(self.root,text=f"Copyright @ {date_} RootTech", font=("times new roman",12,"bold")).pack(side=BOTTOM,fill=X)
        permission.interact_with_database((resource_path('PRMS.db'))) 
        self.update_date_time()
        self.show()
        
        
    def show(self):
        try:
            self.cur.execute("Select pat_name,intervention, amount_paid, date from doctor_patient_records")
            rows=self.cur.fetchall()
            self.Cash_BookTable.delete(*self.Cash_BookTable.get_children())
            for row in rows:
                self.Cash_BookTable.insert('',END,values=row)
        except sqlite3.Error as ex:
            messagebox.showerror("Error",f"Database Error due to : {str(ex)}",parent=self.root)  
            
            
    def search(self):
        try:
            if self.var_search.get()=="":
                messagebox.showerror("Error","Search input is required",parent=self.root)
            else:        
                self.cur.execute("select pat_name,intervention, amount_paid, date from doctor_patient_records where pat_name LIKE '%"+self.var_search.get()+"%'")
                rows=self.cur.fetchall()
                if len(rows)!=0:
                    self.Cash_BookTable.delete(*self.Cash_BookTable.get_children())
                    for row in rows:
                        self.Cash_BookTable.insert('',END,values=row)
                else:
                    messagebox.showerror("Error","No record found!!!",parent=self.root)
        except sqlite3.Error as ex:
            messagebox.showerror("Error",f"Database Error due to : {str(ex)}",parent=self.root)
                      
        
    def sort_treeTable(self,):
        try:
            month = self.month_.get()
            year = self.year_.get()
            if month=='' or year=='':
                messagebox.showerror('Invalid',"Please add a valid month ('00') and year ('0000')")
                return
            else:
                # Convert month and year to a string in the format 'YYYY-MM'
                self.date_str = f"{month}.{year}"

                # Execute SQL query to select records from the specified month and year
                self.cur.execute(f"SELECT pat_name, intervention, amount_paid, date FROM doctor_patient_records WHERE strftime('%m.%Y', substr(date,7,4) || '-' || substr(date,4,2) || '-' || substr(date,1,2)) = '{self.date_str}'")
                rows = self.cur.fetchall()

                # Delete existing records in the Treeview
                self.Cash_BookTable.delete(*self.Cash_BookTable.get_children())

                # Insert the fetched records into the Treeview
                for row in rows:
                    self.Cash_BookTable.insert('', END, values=row)
        except Exception as ex:
            messagebox.showerror("Error", f"Error due to : {str(ex)}", parent=self.root)
   

    def generate_cash_book(self):
        try:
            month = self.month_.get()
            year = self.year_.get()
            if month=='' or year=='':
                messagebox.showerror('Invalid',"Please add a valid month ('00') and year ('0000')")
                return
            else:
                # Convert month and year to a string in the format 'MM.YYYY'
                self.date_str = f"{month}.{year}"
                self.exist=os.path.join(os.getcwd(),resource_path(f'Cash_Book\\cash_book_{self.date_str}.docx'))
                if not os.path.exists(self.exist):
                    self.cashbook(self.date_str)
                else:
                    op=messagebox.askyesno("Error", f"Cash book for {self.date_str} already exist!!!\nDo you want to Regenerate it?", parent=self.root)
                    if op == True:
                        self.cashbook(self.date_str)
        except Exception as ex:
            messagebox.showerror("Error", f"Error due to : {str(ex)}", parent=self.root)
            
    def cashbook(self,date_str):
        # Execute SQL query to select records from the specified month and year
            self.cur.execute(f"SELECT pat_name, intervention, amount_paid, date FROM doctor_patient_records WHERE strftime('%m.%Y', substr(date,7,4) || '-' || substr(date,4,2) || '-' || substr(date,1,2)) = '{self.date_str}'")
            rows = self.cur.fetchall()

            # Create a new Document
            doc = Document()
            header=doc.sections[0].header
            for paragraph in header.paragraphs:
                for run in paragraph.runs:
                    run.clear()
            htable=header.add_table(1,2,Inches(6))
            htab_cells=htable.rows[0].cells
            
            ht0=htab_cells[0].add_paragraph()
            logo=ht0.add_run()
            logo.add_picture(resource_path('images\\emma.png'))
            ht1=htab_cells[1].add_paragraph(f'''\tEMMANUEL DENTAL CLINIC 
            BONAMOUSSADI
            Pr.(Dr) AGBOR MICHEAL ASHU
            Tel: 677 17 01 67/697 12 27 82
            www.emmanueldentalcare.org
            ''')
            ht1.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p_date=doc.add_paragraph(f"{date_str}")
            p_date.alignment=WD_ALIGN_PARAGRAPH.RIGHT
            p_date.add_run(f"\n")
       
            # Add a title to the document
            title=doc.add_heading(f'''\tCASH BOOK
    ''',level=0)
            title.alignment=WD_ALIGN_PARAGRAPH.CENTER

            
            # Add a table to the document
            table = doc.add_table(rows=1, cols=5)

            # Add the column names to the table
            for i, name in enumerate(["Patient Name", "Intervention", "Amount Paid", "Date", "Total"]):
                table.cell(0, i).text = name

            # Initialize a dictionary to store the total amount paid for each patient
            total_amount_paid = {}

            # Initialize a variable to store the total amount paid by all patients
            grand_total = 0

            # Add the fetched records to the table
            for row in rows:
                cells = table.add_row().cells
                for i, value in enumerate(row):
                    cells[i].text = str(value)

                # Add the amount paid to the total for this patient
                pat_name = row[0]
                amount_paid = float(row[2])
                if pat_name in total_amount_paid:
                    total_amount_paid[pat_name] += amount_paid
                else:
                    total_amount_paid[pat_name] = amount_paid

                # Add the amount paid to the grand total
                grand_total += amount_paid

                # Add the total amount paid for this patient to the table
                cells[4].text = str(total_amount_paid[pat_name])

            # Add a row to the table for the grand total
            cells = table.add_row().cells
            cells[0].text = "Grand Total"
            cells[4].text = str(grand_total)
           
            data = doc.add_paragraph()
            data.add_run(f"\n")
            data.add_run(f"Total: ").bold=True
            data.add_run(f"{str(grand_total)} FCFA").bold=True
            data.add_run(f"\n")
            data.add_run(f"\t{num2words(grand_total).upper()} FCFA").italic = True
            data.add_run(f"\n")
            data.add_run(f"\n")
            data.add_run(f"\t\t\tSincerely,").bold=True
            data.add_run(f"\n")
            data.add_run(f"\t\t\tPr (Dr). Agbor Michael Ashu").bold=True

            # Save the document
            doc.save(os.path.join(os.getcwd(),resource_path(f'Cash_Book\\cash_book_{date_str}.docx')))
            messagebox.showinfo("Success", f"Cash Book for {date_str} has been Generated", parent=self.root)

    def print_cashbook(self):
        month = self.month_.get()
        year = self.year_.get()
        self.date_str = f"{month}.{year}"
        file_path=os.path.join(os.getcwd(),resource_path(f'Cash_Book\\cash_book_{self.date_str}.docx'))
        if not os.path.exists(file_path):
            messagebox.showerror("Print error",f"The file {file_path} does not exist.Generate CashBook First",parent=self.root)
        else:
            notification.notify(title="Printing",Message=f"Printing CashBook for {file_path}!!! be patient",timeout=40)
            doc=docx.Document(file_path)
            printer_name=win32print.GetDefaultPrinter()
            win32api.ShellExecute(0,"print",file_path,'d:"%s"'%printer_name,".",0)
    
    def update_date_time(self):
        time_ = time.strftime("%H:%M:%S")
        date_ = time.strftime("%d:%m:%Y")
        self.lbl_clock.configure(text=f"Welcome to SmileScribePro - Professional Patient Records Management System\t\t Date: {str(date_)}\t\t Time: {str(time_)}")
        self.lbl_clock.after(200, self.update_date_time)
        

    def exit(self):
        self.is_running = False
        self.root.destroy()
    
        
if __name__=="__main__":
    root=ctk.CTk()
    obj=cashBookClass(root)
    root.mainloop()