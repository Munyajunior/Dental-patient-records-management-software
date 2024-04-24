from tkinter import*
import tkinter as tk
import customtkinter as ctk
from PIL import Image
from tkinter import ttk,messagebox
import permission
import sqlite3 
import time
import os
import docx
from docx import Document
from docx.enum.text import  WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt,RGBColor
from num2words import num2words
import sys
import win32print
import win32api

def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS2
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)



ctk.set_appearance_mode("system")
ctk.set_default_color_theme("blue")

class ProformaClass(ctk.CTk):
    def __init__(self,root):
        self.root=root
        self.root.state('zoomed')
        # Bind the exit method to the window's close button
        self.root.protocol("WM_DELETE_WINDOW", self.exit)
        screen_width=self.root.winfo_screenwidth()
        screen_height=self.root.winfo_screenheight()
        self.root.geometry(f'{screen_width}x{screen_height}')
        self.root.iconbitmap(resource_path('icon.ico'))
        self.root.title("SmileScribePro")
        self.cart_list=[]
        self.proforma_data_list=[]
        #============================Title======================
        self.icon_title=ctk.CTkImage(dark_image=Image.open(resource_path("images\\logo.png")),
                                    light_image=Image.open(resource_path("images\\logo.png")),size=(50,50))
        title = ctk.CTkLabel(
        self.root,
        text="SmileScribePro",
        image=self.icon_title,
        compound=LEFT,
        font=("sans-serif", 40, "bold"),
        fg_color=("#3498db","#fff"),
        text_color=("#fff","#3498db"),
        anchor='w',
        width=screen_width,  # Set the width to be the screen width
        height=68,padx=20
    ).place(anchor='nw', x=0, y=0)
        #=============Btn_Logout================================
        # Calculate the position of the button based on the screen size
        x_position = screen_width - 100  # Adjust this value as needed
        y_position = screen_height - 735  # Adjust this value as needed
        btn_logout = ctk.CTkButton(self.root,text="Exit"
                                    ,command=self.exit
                                    ,font=("calibri",20,"bold")
                                    ,fg_color=("#273c75","#353b48")
                                    ,hover_color="#e84118"
                                    ,height=50,width=150
                                    ,corner_radius=0
                                    ).place(anchor='e', x=x_position, y=30)
        #======================Clock============================
        self.lbl_clock=ctk.CTkLabel(self.root,text="Welcome to SmileScribePro - Professional Patient Records Management System\t\t Date: DD-MM-YYYY\t\t Time: HH:MM:SS"
                                    ,font=("calibri",18,"bold")
                                    ,fg_color=("#487eb0","#4d636d"),bg_color="#fff"
                                    ,width=screen_width
                                    ,text_color="#fff")
        self.lbl_clock.place(x=0,y=65)
 

           #=================Intervention Main frame================
        InterventionFrame = ctk.CTkFrame(self.root)
        InterventionFrame.place(x=10,y=100) 
        
        pTitle=ctk.CTkLabel(InterventionFrame,text="All Interventions",font=("goudy old style",20,"bold"),width=380).pack(side=TOP,fill=X)
        
        #=================Intervention Search frame================
        #==========Variables================
        self.var_search=StringVar()
        InterventionFrame2=ctk.CTkFrame(self.root)
        InterventionFrame2.place(x=10,y=135) 
        
        lbl_search=ctk.CTkLabel(InterventionFrame2,text="Search Intervention | By Name ",font=("times new roman",15,"bold")).pack(side=TOP,fill=X)
        
        lbl_search=ctk.CTkLabel(self.root,text="Intervention",font=("times new roman",15,"bold")).place(x=10,y=180)
        txt_search=ctk.CTkEntry(self.root,textvariable=self.var_search,font=("times new roman",15),fg_color="lightyellow",width=150,height=22).place(x=100,y=180)
        btn_search=ctk.CTkButton(self.root,text="Search",command=self.search,font=("goudy old style",15),fg_color="#2196f3",width=100,height=25).place(x=290,y=180)
        btn_show_all=ctk.CTkButton(self.root,text="Show All",command=self.show,font=("goudy old style",15),fg_color="#083531",width=100,height=25).place(x=290,y=150)
        
        
        #=====================Intervention Frame============
        InterventionFrame3=ctk.CTkFrame(self.root)
        InterventionFrame3.place(x=10,y=210)
        
        scolly=Scrollbar(InterventionFrame3,orient=VERTICAL)
        scollx=Scrollbar(InterventionFrame3,orient=HORIZONTAL)
        
        self.Intervention_Table=ttk.Treeview(InterventionFrame3,columns=("tp_id","tp_name","tp_code","tp_price"),yscrollcommand=scolly.set,xscrollcommand=scollx.set)
        scollx.pack(side=BOTTOM,fill=X)
        scolly.pack(side=RIGHT,fill=Y)
        scollx.configure(command=self.Intervention_Table.xview)
        scolly.configure(command=self.Intervention_Table.yview)
        
        self.Intervention_Table.heading("tp_id",text="ITV ID")
        self.Intervention_Table.heading("tp_name",text="Intervention") 
        self.Intervention_Table.heading("tp_code",text="Code")
        self.Intervention_Table.heading("tp_price",text="Price")     
                
        self.Intervention_Table["show"] ="headings" 
               
        self.Intervention_Table.column("tp_id",width=90)
        self.Intervention_Table.column("tp_name",width=100)
        self.Intervention_Table.column("tp_code",width=80)
        self.Intervention_Table.column("tp_price",width=100)            
        self.Intervention_Table.pack(fill=BOTH,expand=1)
        self.Intervention_Table.bind("<ButtonRelease-1>",self.get_data)
        
        
        #============Patient Frame==========
        PatientFrame = ctk.CTkFrame(self.root)
        PatientFrame.place(x=10,y=460)
        pTitle=ctk.CTkLabel(PatientFrame,text="All Patients",font=("cambria",17))
        pTitle.pack(side=TOP,fill=X)
        
        scrolly = Scrollbar(PatientFrame,orient=VERTICAL)
        scrollx = Scrollbar(PatientFrame,orient=HORIZONTAL)
        
        self.patient_table =ttk.Treeview(PatientFrame,columns=("pat_id","name","tooth","tp","date"))
        scrollx.pack(side=BOTTOM,fill=X)
        scrolly.pack(side=RIGHT,fill=Y)
        self.patient_table.configure(xscrollcommand=scrollx.set, yscrollcommand=scrolly.set)
        scrollx.config(command=self.patient_table.xview)
        scrolly.config(command=self.patient_table.yview)
        
        self.patient_table.heading("pat_id",text="PAT ID")
        self.patient_table.heading("name",text="Name") 
        self.patient_table.heading("tooth",text="Teeth")  
        self.patient_table.heading("tp",text="Interventions")  
        self.patient_table.heading("date",text="Date")        
        self.patient_table["show"] ="headings" 
               
        self.patient_table.column("pat_id",width=60)
        self.patient_table.column("name",width=100)
        self.patient_table.column("tooth",width=100)
        self.patient_table.column("tp",width=100)
        self.patient_table.column("date",width=10)            
        self.patient_table.pack(fill=BOTH,expand=1)
        self.patient_table.bind("<ButtonRelease-1>",self.get_patient_data)
        

        #====================================Cart Frame===================  
        cart_Frame=ctk.CTkFrame(self.root)
        cart_Frame.place(x=410,y=100)
        self.cartTitle=ctk.CTkLabel(cart_Frame,text="Manage Patient Proforma",font=("goudy old style",17),width=440)
        self.cartTitle.pack(side=TOP,fill=X)
        
        scolly=Scrollbar(cart_Frame,orient=VERTICAL)
        scollx=Scrollbar(cart_Frame,orient=HORIZONTAL)
        
        self.CartTable=ttk.Treeview(cart_Frame,columns=("tp_name","tp_code","tp_price","teeth"),yscrollcommand=scolly.set,xscrollcommand=scollx.set)
        scollx.pack(side=BOTTOM,fill=X)
        scolly.pack(side=RIGHT,fill=Y)
        scollx.config(command=self.CartTable.xview)
        scolly.config(command=self.CartTable.yview)
        self.CartTable.heading("tp_name",text="Intervention") 
        self.CartTable.heading("tp_code",text="Code")
        self.CartTable.heading("tp_price",text="Price")  
        self.CartTable.heading("teeth",text="Teeth")     
                
        self.CartTable["show"] ="headings" 
               
        self.CartTable.column("tp_name",width=100)
        self.CartTable.column("tp_code",width=100)
        self.CartTable.column("tp_price",width=150)
        self.CartTable.column("teeth",width=150)            
        self.CartTable.pack(fill=BOTH,expand=1)
        self.CartTable.bind("<ButtonRelease-1>",self.get_proforma_data)
       
       #=======================Add Cart Widgets Frame===================
        #=====Variable=====
        self.var_tp_name=StringVar()
        self.var_price=StringVar()
        self.var_code=StringVar()
        self.var_tooth=StringVar()
        self.var_pname=StringVar()
        
        Add_cartWidgetsFrame=ctk.CTkFrame(self.root,width=530,height=200)
        Add_cartWidgetsFrame.place(x=410,y=390)
        
        lbl_p_name=ctk.CTkLabel(Add_cartWidgetsFrame,text="Intervention",font=("times new roman",15)).place(x=5,y=5)       
        txt_p_name=ctk.CTkEntry(Add_cartWidgetsFrame,textvariable=self.var_tp_name,font=("times new roman",15),fg_color="lightyellow",width=190,height=22).place(x=5,y=35)
        
        lbl_p_price=ctk.CTkLabel(Add_cartWidgetsFrame,text="Code",font=("times new roman",15)).place(x=230,y=5)       
        txt_p_price=ctk.CTkEntry(Add_cartWidgetsFrame,textvariable=self.var_code,font=("times new roman",15),fg_color="lightyellow",state='readonly',width=150,height=22).place(x=230,y=35)
        
        lbl_p_qty=ctk.CTkLabel(Add_cartWidgetsFrame,text="Price",font=("times new roman",15)).place(x=390,y=5)       
        txt_p_qty=ctk.CTkEntry(Add_cartWidgetsFrame,textvariable=self.var_price,font=("times new roman",15),fg_color="lightyellow",state='readonly',width=130,height=22).place(x=390,y=35)
        
        txt_p_qty=ctk.CTkEntry(Add_cartWidgetsFrame,textvariable=self.var_pname,font=("times new roman",15),fg_color="lightyellow",state='readonly',width=190,height=22).place(x=5,y=120)
        
        
        lbl_tooth=ctk.CTkLabel(Add_cartWidgetsFrame,text="Tooth",font=("times new roman",15)).place(x=5,y=60)
        self.tooth_type = ttk.Combobox(Add_cartWidgetsFrame,textvariable=self.var_tooth,values=["Select","Primary","Permanent"],state='readonly',justify=CENTER)
        self.tooth_type.set("Select")
        self.tooth_type.bind("<<ComboboxSelected>>", self.on_tooth_type_selected)
        self.tooth_type.place(x=5,y=85,width=200)

        
        self.tooth_selection = tk.Listbox(Add_cartWidgetsFrame,selectmode="multiple",exportselection=0, state='disabled')  # Initially disable the listbox
        self.tooth_selection.bind("<<ListboxSelect>>", self.on_tooth_selected)  # Bind the selection event
        self.tooth_selection.place(x=205,y=40,width=80,height=100)
        
        
        btn_clear_cart=ctk.CTkButton(Add_cartWidgetsFrame,text="Clear",command=self.clear_cart,font=("times new roman",15,"bold"),fg_color="#eb4d4b",width=150,height=30).place(x=340,y=110)      
        btn_add_cart=ctk.CTkButton(Add_cartWidgetsFrame,text="Add Data",command=self.add_update_cart,font=("times new roman",15,"bold"),fg_color="orange",width=180,height=30).place(x=340,y=70)   
       
        lbl_note=ctk.CTkLabel(Add_cartWidgetsFrame,text=f"Note:'Enter 0 Intervention to remove Intervention from the Proforma Manager'",font=("goudy old style",15),anchor='se',text_color="red",justify=CENTER).place(x=10,y=150)
        
       
        #================Proforma Area=====
        billFrame=ctk.CTkFrame(self.root)
        billFrame.place(x=960,y=100)
        
        BTitle=ctk.CTkLabel(billFrame,text="HAVE FUN PLAYING GUESS GAME",font=("goudy old style",15,"bold"),fg_color="#f44336",text_color="white").pack(side=TOP,fill=X)
        scrolly=Scrollbar(billFrame,orient=VERTICAL)
        scrollx=Scrollbar(billFrame,orient=HORIZONTAL)
        scrolly.pack(side=RIGHT,fill=Y)
        scrollx.pack(side=BOTTOM,fill=X)
        
        
        self.txt_game_area=ctk.CTkTextbox(billFrame,font=("courier new",12,"bold"),text_color="#2c3e50",bg_color="#ecf0f1",width=390,height=420)
        self.txt_game_area.insert(tk.END, "Guess a number between 1 and 100\n")
        self.txt_game_area.configure(yscrollcommand=scrolly.set,xscrollcommand=scrollx.set)
        scrollx.config(command=self.txt_game_area.xview)
        scrolly.config(command=self.txt_game_area.yview)
        self.txt_game_area.pack(fill=BOTH,expand=1)
       
        
        
        #+++++++++++++++++++++++++bILLING Button==============
        billMenuFrame=ctk.CTkFrame(self.root,width=400,height=150)
        billMenuFrame.place(x=960,y=570)
 
        self.lbl_amnt=ctk.CTkLabel(billMenuFrame,text='Proforma\nAmount [0]',font=("goudy old style",15,"bold"),fg_color="#3f51b5",text_color="white",width=153,height=60)
        self.lbl_amnt.place(x=2,y=5)
        
        self.lbl_netpay=ctk.CTkLabel(billMenuFrame,text='Proforma \nNet Pay [0]',font=("goudy old style",15,"bold"),fg_color="#abc34a",text_color="white",width=100,height=60)
        self.lbl_netpay.place(x=157,y=5)
        
        btn_reset_game=ctk.CTkButton(billMenuFrame,text='Play Game',command=self.check_guess,font=("goudy old style",15,"bold"),fg_color="#607d8b",text_color="white",width=140,height=60)
        btn_reset_game.place(x=260,y=5)
        
        btn_print=ctk.CTkButton(billMenuFrame,text='Generate\nProforma',command=self.generate_proforma,font=("goudy old style",15,"bold"),fg_color="lightgreen",text_color="white",width=120,height=60)
        btn_print.place(x=2,y=70)
        
        btn_clear_all=ctk.CTkButton(billMenuFrame,text='View\nProforma',command=self.view_proforma,font=("goudy old style",15,"bold"),fg_color="gray",text_color="white",width=120,height=60)
        btn_clear_all.place(x=124,y=70)
        
        btn_generate=ctk.CTkButton(billMenuFrame,text='Print\nProforma',command=self.print_proforma,font=("goudy old style",15,"bold"),fg_color="#009688",text_color="white",width=155,height=60)
        btn_generate.place(x=246,y=70)
        
        #==================Footer================
        date_=time.strftime("%Y")
        lbl_footer=ctk.CTkLabel(self.root,text=f"Copyright @ {date_} RootTech", font=("times new roman",12,"bold")).pack(side=BOTTOM,fill=X)
        
        self.update_date_time()
        self.show()
        self.show_data('patient', clear=True)  # Clear the treeview and show data from 'patient'
        self.show_data('archives', clear=False)  # Don't clear the treeview, just add data from 'archives'

        
        
        
        
    #==========================ALL FUNCTIONS====================
 
    def on_tooth_type_selected(self, event):
        self.tooth_selection['state'] = 'normal'  # Enable the listbox when a selection is made
        self.tooth_selection.delete(0, tk.END)
        tooth_type = self.tooth_type.get()
        if tooth_type == "Primary":
            self.teeth = ["#51", "#52", "#53","#54","#55","#61","#61","#62","#63",
                        "#64","#65","#71","#72","#73","#74","#75","#81","#82","#83","#84","#85"]  
        else:
            self.teeth = ["#18","#17","#16","#15","#14","#13","#12","#11",
                        "#21","#22","#23","#24","#25","#26","#27","#28","#48",
                        "#47","#46","#45","#44","#43","#42","#41","#31","#32","#33",
                        "#34","#35","#36","#37","#38"] 
        for tooth in self.teeth:
            self.tooth_selection.place()
            self.tooth_selection.insert(tk.END, tooth)
            
    def on_tooth_selected(self, event):  # New function to handle selection event
        selected_teeth = [self.tooth_selection.get(idx) for idx in self.tooth_selection.curselection()]
        self.var_tooth.set(", ".join(selected_teeth))
     
    def update_date_time(self):   
        time_=time.strftime("%H:%M:%S")
        date_=time.strftime("%d:%m:%Y")
        self.lbl_clock.configure(text=f"Welcome to SmileScribe Professional Patients Record Management System\t\t Date: {str(date_)}\t\t Time: {str(time_)}")
        self.lbl_clock.after(200,self.update_date_time)
        
    def show(self):
        permission.interact_with_database((resource_path('PRMS.db')))
        con=sqlite3.connect(database=resource_path(r'PRMS.db'))
        cur=con.cursor()
        try:
            cur.execute("Select tp_id,tp_name,tp_code,tp_price from treatment")
            rows=cur.fetchall()
            self.Intervention_Table.delete(*self.Intervention_Table.get_children())
            for row in rows:
                self.Intervention_Table.insert('',END,values=row)
        except Exception as ex:
            messagebox.showerror("Error",f"Error due to : {str(ex)}",parent=self.root)

    def show_data(self, table_name, clear=True):
        permission.interact_with_database((resource_path('PRMS.db')))
        con=sqlite3.connect(database=resource_path(r'PRMS.db'))
        cur=con.cursor()
        try:
            cur.execute(f"Select pat_id,name,tooth,tp,date from {table_name}")
            rows=cur.fetchall()
            if clear:
                self.patient_table.delete(*self.patient_table.get_children())
            for row in rows:
                self.patient_table.insert('',END,values=row)
        except Exception as ex:
            messagebox.showerror("Error",f"Error due to : {str(ex)}",parent=self.root)

    
    
    '''def show_patient(self):
        permission.interact_with_database((resource_path('PRMS.db')))
        con=sqlite3.connect(database=resource_path(r'PRMS.db'))
        cur=con.cursor()
        try:
            cur.execute("Select pat_id,name,tooth,tp,date from patient")
            rows=cur.fetchall()
            self.patient_table.delete(*self.patient_table.get_children())
            for row in rows:
                self.patient_table.insert('',END,values=row)
        except Exception as ex:
            messagebox.showerror("Error",f"Error due to : {str(ex)}",parent=self.root)      
    
    def show_archive(self):
        permission.interact_with_database((resource_path('PRMS.db')))
        con=sqlite3.connect(database=resource_path(r'PRMS.db'))
        cur=con.cursor()
        try:
            cur.execute("Select pat_id,name,tooth,tp,date from archives")
            rows=cur.fetchall()
            self.patient_table.delete(*self.patient_table.get_children())
            for row in rows:
                self.patient_table.insert('',END,values=row)
        except Exception as ex:
            messagebox.showerror("Error",f"Error due to : {str(ex)}",parent=self.root)            
    '''          
    def search(self):
        permission.interact_with_database((resource_path('PRMS.db'))) 
        con=sqlite3.connect(database=resource_path(r'PRMS.db'))
        cur=con.cursor()
        try:
            if self.var_search.get()=="":
                messagebox.showerror("Error","Search input is required",parent=self.root)
            else:        
                cur.execute("select tp_id,tp_name,tp_code,tp_price from treatment where tp_name LIKE '%"+self.var_search.get()+"%'")
                rows=cur.fetchall()
                if len(rows)!=0:
                    self.Intervention_Table.delete(*self.Intervention_Table.get_children())
                    for row in rows:
                        self.Intervention_Table.insert('',END,values=row)
                else:
                    messagebox.showerror("Error","No record found!!!",parent=self.root)
        except Exception as ex:
            messagebox.showerror("Error",f"Error due to : {str(ex)}",parent=self.root)        
            
    def get_data(self,ev):
        f=self.Intervention_Table.focus()
        content=(self.Intervention_Table.item(f))  
        row=content['values']
        self.var_tp_name.set(row[1])
        self.var_code.set(row[2])
        self.var_price.set(row[3])
    
    def get_patient_data(self,ev):
        f=self.patient_table.focus()
        content=(self.patient_table.item(f))
        row=content["values"]
        self.var_pname.set(row[1])
                   
    def get_proforma_data(self,ev):
        f=self.CartTable.focus()
        content=(self.CartTable.item(f))  
        row=content['values']
        self.var_tp_name.set(row[0])
        self.var_code.set(row[1])
        self.var_price.set(row[2]) 
     
    def add_update_cart(self):
        if  self.var_tp_name.get()=='':
            messagebox.showerror('Error',"Please select Intervention from the List",parent=self.root)            
        elif self.var_price.get()=='':
            messagebox.showerror('Error',"Price is Required",parent=self.root)
        else:
            price_cal=self.var_price.get()
            cart_data=[self.var_tp_name.get(),self.var_code.get(),price_cal,self.var_tooth.get()]
            
            #=========update cart===============
            present='no'
            index_=0
            for row in self.cart_list:
                if self.var_code.get()==row[1]:
                    present='yes'
                    break
                index_+=1
            if present=='yes':
                op=messagebox.askyesno('Confirm',"Intervention already present do you want to \nAdd or Update Intervention in the Patient Proforma Manager?",parent=self.root)
                if op==True:
                    if self.var_tp_name.get()=='0' :
                        self.cart_list.pop(index_)    
                        messagebox.showinfo("Success","Intervention has been Successfully deleted",parent=self.root)
                    elif self.cart_list[index_][3]==self.var_tooth.get():
                         messagebox.showerror("Error",f'''{self.var_tp_name.get()} cannot be done twice on the same {self.var_tooth.get()}''',parent=self.root)
                    elif self.cart_list[index_][3] == '':
                        self.cart_list[index_][3]=self.var_tooth.get()
                        messagebox.showinfo("Update Successful","Intervention has been Successfully Updated",parent=self.root)
                    elif self.cart_list[index_][3] == self.var_tooth.get():
                        self.cart_list[index_][3]=self.var_tooth.get()
                        messagebox.showinfo("Update Successful","Intervention has been Successfully Updated",parent=self.root)
                    elif self.cart_list[index_][3] != self.var_tooth.get():
                        self.cart_list.append(cart_data)
                        messagebox.showinfo("Added Successful","Intervention has been Successfully Added",parent=self.root)
                
            else:
                self.cart_list.append(cart_data)
            self.show_cart()
            self.update_proforma_bill()
            
     
            
            
    def show_cart(self):
        try:
            self.CartTable.delete(*self.CartTable.get_children())
            for row in self.cart_list:
                self.CartTable.insert('',END,values=row)
        except Exception as ex:
            messagebox.showerror("Error",f"Error due to : {str(ex)}",parent=self.root)       
    
    def get_cartTable_data(self):
        # Get all items in the Treeview
        for item in self.CartTable.get_children():
            # Get the item's values
            item_values = self.CartTable.item(item)['values']
            # Append the item's values to the cart list
            self.proforma_data_list.append(item_values)
    
    def add_data_to_document(self, document, data_list, name):
        total = sum(int(item[2]) for item in data_list if item[2] != '')
        date = time.strftime('%d-%m-%Y')
        word = num2words(total)
        # Check if the style 'Title' already exists, if not then create it
        if 'Title' not in document.styles:
            style = document.styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(24)
            # Set the font color to dark blue (25% darker)
            font.color.rgb = RGBColor(23,54,93)
        # Now you can add a heading with the 'Title' style
        title=document.add_heading(f"PROFORMA {name.upper()}",level=0)
        title.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p_date=document.add_paragraph(f"{date}")
        p_date.alignment=WD_ALIGN_PARAGRAPH.RIGHT
        p_date.add_run(f"\n")
        title.add_run(f"\n")
        for item in data_list:
            intervention, code, price, tooth = item
            p_data=document.add_paragraph()
            p_data.add_run(f"{intervention}\t{code}\t{tooth}\t\t\t\t\t\t{str(price)}")
            p_data.add_run(f"\n")
            p_data.add_run(f"\n")
        p_data.add_run(f"Total: ").bold=True
        p_data.add_run(f"{str(total)} FCFA").bold=True
        p_data.add_run(f"\n")
        p_data.add_run(f"\t{word.upper()} FCFA").italic = True
        p_data.add_run(f"\n")
        p_data.add_run(f"\n")
        p_data.add_run(f"\t\t\tSincerely,").bold=True
        p_data.add_run(f"\n")
        p_data.add_run(f"\t\t\tPr (Dr). Agbor Michael Ashu").bold=True
        messagebox.showinfo("Success",f"{name.upper()} Proforma Has Been Generated and saved",parent=self.root)
        
        
        
    def generate_proforma(self):
        try:
            file_path=os.path.join(os.getcwd(), resource_path(f'proforma\\proforma_{self.var_pname.get()}.docx'))
            if os.path.exists(file_path):
                op=messagebox.askyesno("Confirm",f"{self.var_pname.get()} proforma already exist. \nAre you certain you want to Generate a new proforma for this patient",parent=self.root)
                if op==False:
                    return
            if self.var_pname.get()=='':
                messagebox.showerror("Error","Please select Patient whose proforma you are creating from all patients",parent=self.root)
                return
            doc = Document(os.path.join(os.getcwd(), resource_path('template.docx')))
            self.get_cartTable_data()
            self.add_data_to_document(doc,self.proforma_data_list,self.var_pname.get())
            doc.save(file_path)
        except Exception as ex:
            messagebox.showerror("Error",f'Error due to: {str(ex)}',parent=self.root)
      
            
    def view_proforma(self):
        file_path=os.path.join(os.getcwd(),resource_path(f'proforma\\proforma_{self.var_pname.get()}.docx'))
        if not os.path.exists(file_path):
            messagebox.showerror("View Error","Generate Patient Proforma First before viewing",parent=self.root) 
        else:
             os.startfile(file_path)  
                    
     
    def print_proforma(self):
            file_path=os.path.join(os.getcwd(), resource_path(f'proforma\\proforma_{self.var_pname.get()}.docx'))
            if not os.path.exists(file_path):
                messagebox.showerror("Print error",f"The file {file_path} does not exist.Generate Patient Proforma First",parent=self.root)
            else:
                doc=docx.Document(file_path)
                printer_name=win32print.GetDefaultPrinter()
                win32api.ShellExecute(0,"print",file_path,'d:"%s"'%printer_name,".",0) 
             
    def update_proforma_bill(self):
        self.bill_amnt=0
        self.net_pay=0
        self.discount=0
        for row in self.cart_list:
            self.bill_amnt=self.bill_amnt+int(row[2])#price is row 2 in cart list
         
        self.net_pay=self.bill_amnt-self.discount
        self.lbl_amnt.configure(text=f"Proforma\nAmount [{str(self.bill_amnt)}]")     
        self.lbl_netpay.configure(text=f"Net Pay \n[{str(self.net_pay)}]")
             

    def check_guess(self):
        last_line = self.txt_game_area.get('end-4c', 'end-1c').strip()
        if last_line.isdigit():
            guess = int(last_line)
            if 1 <= guess <= 100:
                self.guesses += 1
                if guess < self.answer:
                    self.txt_game_area.insert(END, "\nToo low! Try again.\n")
                elif guess > self.answer:
                    self.txt_game_area.insert(END, "\nToo high! Try again.\n")
                else:
                    self.txt_game_area.delete('1.0',END)
                    self.txt_game_area.insert(END, f"\nCongratulations! You found the number in {self.guesses} guesses.\n")
                    self.txt_game_area.insert(END, "Play Again!!! Guess a number between 1 and 100\n")
            else:
                self.txt_game_area.insert(END, "\nInvalid input! Please enter a number between 1 and 100.\n")
        else:
            self.txt_game_area.insert(END, "\nInvalid input! Please enter a number between 1 and 100.\n")
                    
    def clear_cart(self):
        self.var_tp_name.set('')
        self.var_price.set('')
        self.var_code.set('')   
        self.var_tooth.set('Select')
              
    def exit(self):
        self.is_running = False
        self.root.destroy() 
    
if __name__=="__main__":
    root=ctk.CTk()
    obj=ProformaClass(root)
    root.mainloop()