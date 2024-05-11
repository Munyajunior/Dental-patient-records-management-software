from tkinter import*
import customtkinter as ctk
from PIL import Image
from tkinter import ttk,messagebox
import sqlite3
import time
import docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import win32print
import win32api
import docx
from docx import Document
from docx.enum.text import  WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt,RGBColor
from num2words import num2words
from datetime import datetime
import os
import permission
import sys


def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS2
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)



ctk.set_appearance_mode("system")
ctk.set_default_color_theme("blue")


class quoteClass(ctk.CTk):
    def __init__(self,root):
        self.root=root
        self.root.state('zoomed')
        # Bind the exit method to the window's close button
        self.root.protocol("WM_DELETE_WINDOW", self.exit)
        screen_width=self.root.winfo_screenwidth()
        screen_height=self.root.winfo_screenheight()
        self.root.geometry(f'{screen_width}x{screen_height}')
        self.root.iconbitmap(resource_path('icon.ico'))
        self.root.title("SmileScribePro")
        self.quote_list=[]
        self.quote_data_list=[]
        permission.interact_with_database((resource_path('PRMS.db'))) 
        self.con=sqlite3.connect(database=resource_path(r'PRMS.db'))
        self.cur=self.con.cursor()
               
        #============================Title======================
        self.icon_title=ctk.CTkImage(dark_image=Image.open(resource_path("images\\logo.png")),
                                    light_image=Image.open(resource_path("images\\logo.png")),size=(50,50))
        title = ctk.CTkLabel(
        self.root,
        text="SmileScribePro",
        image=self.icon_title,
        compound=LEFT,
        font=("sans-serif", 40, "bold"),
        fg_color=("#3498db","#fff"),
        text_color=("#fff","#3498db"),
        anchor='w',
        width=screen_width,  # Set the width to be the screen width
        height=68,padx=20
    ).place(anchor='nw', x=0, y=0)
        #=============Btn_Logout================================
        # Calculate the position of the button based on the screen size
        x_position = screen_width - 100  # Adjust this value as needed
        y_position = screen_height - 735  # Adjust this value as needed
        btn_logout = ctk.CTkButton(self.root,text="Exit"
                                    ,command=self.exit
                                    ,font=("calibri",20,"bold")
                                    ,fg_color=("#273c75","#353b48")
                                    ,hover_color="#e84118"
                                    ,height=50,width=150
                                    ,corner_radius=0
                                    ).place(anchor='e', x=x_position, y=30)
        #======================Clock============================
        self.lbl_clock=ctk.CTkLabel(self.root,text="Welcome to SmileScribePro - Professional Patient Records Management System\t\t Date: DD-MM-YYYY\t\t Time: HH:MM:SS"
                                    ,font=("calibri",18,"bold")
                                    ,fg_color=("#487eb0","#4d636d"),bg_color="#fff"
                                    ,width=screen_width
                                    ,text_color="#fff")
        self.lbl_clock.place(x=0,y=65)

        #=======================Content=======================
        #===================title=========
        title=ctk.CTkLabel(self.root,text="Doctor Quotes",font=("goudy old style",20,"bold"),fg_color="#0f4d7d",text_color="#fff",width=1000).place(x=170,y=100)
       
       #=================Doctor Records Main frame================
        DoctorFrame = ctk.CTkFrame(self.root)
        DoctorFrame.place(x=60,y=130) 

        pTitle=ctk.CTkLabel(DoctorFrame,text="All Doctor Records",font=("arial",20,"bold"),text_color="#fff",fg_color="#262626",width=380).pack(side=TOP,fill=X)

        #=================Doctor Record Search frame================
        
        #==========Variables================
        self.var_search=ctk.StringVar()

        lbl_search=ctk.CTkLabel(self.root,text="Search Doctor Records | By Name ",font=("times new roman",15,"bold")).place(x=60,y=170)

        lbl_search=ctk.CTkLabel(self.root,text="Records",font=("times new roman",15,"bold")).place(x=60,y=200)
        txt_search=ctk.CTkEntry(self.root,textvariable=self.var_search,font=("times new roman",15),fg_color="lightyellow",width=150,height=22).place(x=150,y=200)
        btn_search=ctk.CTkButton(self.root,text="Search",command=self.search,font=("arial",15),fg_color="#2196f3",width=100,height=25).place(x=350,y=200)
        btn_show_all=ctk.CTkButton(self.root,text="Show All",command=self.show,font=("arial",15),fg_color="#083531",width=100,height=25).place(x=350,y=170)

        #=====================Doctor Records Frame============
        DoctorRecordFrame=ctk.CTkFrame(self.root)
        DoctorRecordFrame.place(x=10,y=240)

        scolly=Scrollbar(DoctorRecordFrame,orient=VERTICAL)
        scollx=Scrollbar(DoctorRecordFrame,orient=HORIZONTAL)
        
        def treeview_sort_column(tv, col, reverse):            
            if col == 'DATE':
                l = [(datetime.strptime(tv.set(k, col), "%d.%m.%Y"), k) for k in tv.get_children('')]
            else:
                l = [(tv.set(k, col), k) for k in tv.get_children('')]
            l.sort(reverse=reverse)
            for index, (val, k) in enumerate(l):
                tv.move(k, '', index)
            tv.heading(col, command=lambda _col=col: treeview_sort_column(tv, _col, not reverse))

        self.Doc_Record_Table=ttk.Treeview(DoctorRecordFrame,columns=("doc_id", "doc_name","pat_name","intervention", "amount_paid", "date"),yscrollcommand=scolly.set,xscrollcommand=scollx.set)
        scolly.pack(side=RIGHT,fill=Y)
        scollx.pack(side=BOTTOM,fill=X)
        scolly.config(command=self.Doc_Record_Table.yview)
        scollx.config(command=self.Doc_Record_Table.xview)

        columns = {"doc_id":"DOC_ID", "doc_name":"DOC_NAME","pat_name":"PATIENT_NAME","intervention":"INTERVENTION", "amount_paid":"AMOUNT_PAID", "date":"DATE"}
        
        for k,v in columns.items():
            self.Doc_Record_Table.heading(k, text=v, command=lambda _col=k: treeview_sort_column(self.Doc_Record_Table, _col, False))

        self.Doc_Record_Table["show"] ="headings" 
        self.Doc_Record_Table.column("doc_id",width=60)
        self.Doc_Record_Table.column("doc_name",width=100)
        self.Doc_Record_Table.column("pat_name",width=100)
        self.Doc_Record_Table.column("intervention",width=100)  
        self.Doc_Record_Table.column("amount_paid",width=100)
        self.Doc_Record_Table.column("date",width=100)          
        self.Doc_Record_Table.pack(fill=BOTH,expand=1)
        self.Doc_Record_Table.bind("<ButtonRelease-1>",self.get_data)

        #=======================Add Quote Widgets Frame===================
        #=====Variable=====
        self.var_doc_id=StringVar()
        self.var_doc_name=StringVar()
        self.var_pat_name=StringVar()
        self.var_interv=StringVar()
        self.var_amt_paid=StringVar()
        self.var_date=StringVar()
        
        Add_QuoteWidgetFrame=ctk.CTkFrame(self.root,width=580,height=200)
        Add_QuoteWidgetFrame.place(x=10,y=490)
        
        lbl_p_name=ctk.CTkLabel(Add_QuoteWidgetFrame,text="Doctor Name",font=("times new roman",15)).place(x=5,y=5)       
        txt_p_name=ctk.CTkEntry(Add_QuoteWidgetFrame,textvariable=self.var_doc_name,font=("times new roman",15)).place(x=5,y=35)
        
        lbl_p_price=ctk.CTkLabel(Add_QuoteWidgetFrame,text="Patient Name",font=("times new roman",15)).place(x=200,y=5)       
        txt_p_price=ctk.CTkEntry(Add_QuoteWidgetFrame,textvariable=self.var_pat_name,font=("times new roman",15),state='readonly').place(x=200,y=35)
        
        lbl_p_qty=ctk.CTkLabel(Add_QuoteWidgetFrame,text="Intervention",font=("times new roman",15)).place(x=380,y=5)       
        txt_p_qty=ctk.CTkEntry(Add_QuoteWidgetFrame,textvariable=self.var_interv,font=("times new roman",15),state='readonly').place(x=380,y=35)
        
        lbl_p_qty=ctk.CTkLabel(Add_QuoteWidgetFrame,text="Total Bill",font=("times new roman",15)).place(x=5,y=70)       
        txt_p_qty=ctk.CTkEntry(Add_QuoteWidgetFrame,textvariable=self.var_amt_paid,font=("times new roman",15),state='readonly').place(x=5,y=100)
        
        btn_clear_cart=ctk.CTkButton(Add_QuoteWidgetFrame,text="Clear",command=self.clear_quote,font=("times new roman",15,"bold"),fg_color="#009688").place(x=200,y=70)      
        btn_add_cart=ctk.CTkButton(Add_QuoteWidgetFrame,text="Add | Update Cart",command=self.add_update,font=("times new roman",15,"bold"),fg_color="orange").place(x=380,y=70)      
 
        lbl_note=ctk.CTkLabel(Add_QuoteWidgetFrame,text="Note:'Enter 0 Doctor Name to remove Intervention from the Bill'",font=("goudy old style",18),text_color="red").place(x=5,y=140)
        
        #=================================Manage Quotes======================
        
        QuoteFrame=ctk.CTkFrame(self.root)
        QuoteFrame.place(x=700,y=240)
        self.quoteTitle=ctk.CTkLabel(QuoteFrame,text="Manage Doctor Quotes",font=("Adobe Caslon Pro",17,))
        self.quoteTitle.pack(side=TOP,fill=X)

        scolly=Scrollbar(QuoteFrame,orient=VERTICAL)
        scollx=Scrollbar(QuoteFrame,orient=HORIZONTAL)

        self.Quote_Table=ttk.Treeview(QuoteFrame,columns=("doc_id", "doc_name","pat_name","intervention", "amount_paid", "date"),yscrollcommand=scolly.set,xscrollcommand=scollx.set)
        scollx.pack(side=BOTTOM,fill=X)
        scolly.pack(side=RIGHT,fill=Y)
        scollx.config(command=self.Quote_Table.xview)
        scolly.config(command=self.Quote_Table.yview)

        self.Quote_Table.heading("doc_id",text="Doctor ID")
        self.Quote_Table.heading("doc_name",text="Doctor Name") 
        self.Quote_Table.heading("pat_name",text="Patient Name")
        self.Quote_Table.heading("intervention",text="Intervention")  
        self.Quote_Table.heading("amount_paid",text="Total Bill") 
        self.Quote_Table.heading("date",text="Date") 
        self.Quote_Table["show"] ="headings" 
        self.Quote_Table.column("doc_id",width=60)
        self.Quote_Table.column("doc_name",width=100)
        self.Quote_Table.column("pat_name",width=100)
        self.Quote_Table.column("intervention",width=100)  
        self.Quote_Table.column("amount_paid",width=100)
        self.Quote_Table.column("date",width=100)          
        self.Quote_Table.pack(fill=BOTH,expand=1)
        self.Quote_Table.bind("<ButtonRelease-1>")
        
        
        
        
        #+++++++++++++++++++++++++Quotes Button==============
        QuoteMenuFrame=ctk.CTkFrame(self.root,width=580,height=180)
        QuoteMenuFrame.place(x=700,y=510)
 
        self.lbl_amnt=ctk.CTkLabel(QuoteMenuFrame,text='Quote\n[0]',font=("goudy old style",15,"bold"),fg_color="#3f51b5",text_color="white",width=153,height=70)
        self.lbl_amnt.place(x=100,y=5)
        
        self.lbl_net_pay=ctk.CTkLabel(QuoteMenuFrame,text='Net Quote\n[0]',font=("goudy old style",15,"bold"),fg_color="#607d8b",text_color="white",width=145,height=70)
        self.lbl_net_pay.place(x=300,y=5)
        
        btn_print=ctk.CTkButton(QuoteMenuFrame,text='Generate Quote',command=self.generate_quote,font=("goudy old style",15,"bold"),fg_color="#6ab04c",text_color="white",width=140,height=50)
        btn_print.place(x=2,y=80)
        
        btn_view=ctk.CTkButton(QuoteMenuFrame,text='View Quote',command=self.view_quote,font=("goudy old style",15,"bold"),fg_color="gray",text_color="white",width=130,height=50)
        btn_view.place(x=150,y=80)
        
        btn_generate=ctk.CTkButton(QuoteMenuFrame,text='Print Quote',command=self.print_quote,font=("goudy old style",15,"bold"),fg_color="#009688",text_color="white",width=130,height=50)
        btn_generate.place(x=290,y=80)
        
        btn_clear_all=ctk.CTkButton(QuoteMenuFrame,text='Clear All',command=self.clear_all,font=("goudy old style",15,"bold"),fg_color="#30336b",text_color="white",width=130,height=50)
        btn_clear_all.place(x=430,y=80)
        
        #==================Footer================
        date_=time.strftime("%Y")
        lbl_footer=ctk.CTkLabel(self.root,text=f"Copyright @ {date_} RootTech", font=("times new roman",12,"bold")).pack(side=BOTTOM,fill=X)
        
        permission.interact_with_database((resource_path('PRMS.db'))) 
        self.show()
        self.update_date_time()
        
        #=====================All functions=================
    
    def show(self):
        try:
            self.cur.execute("Select doc_id, doc_name,pat_name,intervention, amount_paid, date from doctor_patient_records")
            rows=self.cur.fetchall()
            self.Doc_Record_Table.delete(*self.Doc_Record_Table.get_children())
            for row in rows:
                self.Doc_Record_Table.insert('',END,values=row)
        except Exception as ex:
            messagebox.showerror("Error",f"Error due to : {str(ex)}",parent=self.root)
    
    def search(self):
        try:
            if self.var_search.get()=="":
                messagebox.showerror("Error","Search input is required",parent=self.root)
            else:        
                self.cur.execute("select doc_id, doc_name,pat_name,intervention, amount_paid, date from doctor_patient_records where doc_name LIKE '%"+self.var_search.get()+"%'")
                rows=self.cur.fetchall()
                if len(rows)!=0:
                    self.Doc_Record_Table.delete(*self.Doc_Record_Table.get_children())
                    for row in rows:
                        self.Doc_Record_Table.insert('',END,values=row)
                        
                else:
                    messagebox.showerror("Error","No record found!!!",parent=self.root)
        except Exception as ex:
            messagebox.showerror("Error",f"Error due to : {str(ex)}",parent=self.root)
            
            
    def get_data(self,ev):
        f=self.Doc_Record_Table.focus()
        content=(self.Doc_Record_Table.item(f))  
        row=content['values']
        self.var_doc_id.set(row[0])
        self.var_doc_name.set(row[1])
        self.var_pat_name.set(row[2])
        self.var_interv.set(row[3])
        self.var_amt_paid.set(row[4])
        self.var_date.set(row[5])
        
        
    def add_update(self):
        if  self.var_doc_id.get()=='':
            messagebox.showerror('Error',"Please select Doctor Record from the Table",parent=self.root) 
        else:
            price_cal=self.var_amt_paid.get()
            quote_data=[self.var_doc_id.get(),self.var_doc_name.get(),self.var_pat_name.get(),self.var_interv.get(),price_cal,self.var_date.get()]
            
            #=========update quote===============
            present='no'
            index_=0
            for row in self.quote_list:
                if self.var_doc_id.get()==row[0]:
                    present='yes'
                    break
                index_+=1
            if present=='yes':
                op=messagebox.askyesno('Confirm',f"Dr.{self.var_doc_name.get()} already present\nDo you want to Add | Remove from the Quote List",parent=self.root)
                if op==True:
                    if self.var_doc_name.get()=="0":
                        self.quote_list.pop(index_)
                    else:
                        self.quote_list.append(quote_data)
            else:                   
                self.quote_list.append(quote_data)
                
            self.show_quote()
            self.quote_updates()  
            
    def quote_updates(self):
        self.bill_amnt=0
        self.net_pay=0
        self.discount=0
        for row in self.quote_list:
            self.bill_amnt=self.bill_amnt+(int(row[4]))#price is row 5 in cart list
        self.discount=(self.bill_amnt*0)/100   
        self.net_pay=self.bill_amnt-self.discount
        self.lbl_amnt.configure(text=f'Quotes(XAF)\n{str(self.bill_amnt)}')
        self.lbl_net_pay.configure(text=f'Net Quote(XAF)\n{str(self.net_pay)}')
        self.quoteTitle.configure(text=f"Dr.{self.var_doc_name.get()}'s Quote \t Total Quotes: [{str(len(self.quote_list))}]")
        
        
            
    def show_quote(self):
        try:
            self.Quote_Table.delete(*self.Quote_Table.get_children())
            for row in self.quote_list:
                self.Quote_Table.insert('',END,values=row)
        except Exception as ex:
            messagebox.showerror("Error",f"Error due to : {str(ex)}",parent=self.root)  
    
    def clear_quote(self):
        self.var_doc_id.set('')
        self.var_doc_name.set('')
        self.var_pat_name.set('')
        self.var_interv.set('')
        self.var_amt_paid.set('')
        self.var_date.set('')
        
        
    def get_quoteTable(self):
        # Get all items in the Treeview
        for item in self.Quote_Table.get_children():
            # Get the item's values
            item_values = self.Quote_Table.item(item)['values']
            # Append the item's values to the cart list
            self.quote_data_list.append(item_values)
        
    def add_quoting_data_to_document(self,document, data_list, name,date):
        total = sum(int(item[4]) for item in data_list if item[4] != '')
        word = num2words(total)
        # Check if the style 'Title' already exists, if not then create it
        if 'Title' not in document.styles:
            style = document.styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(24)
            # Set the font color to dark blue (25% darker)
            font.color.rgb = RGBColor(23,54,93)
        # Now you can add a heading with the 'Title' style
        title=document.add_heading(f"Dr.{name.upper()}'s MONTHLY QUOTE",level=0)
        title.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p_date=document.add_paragraph(f"{date}")
        p_date.alignment=WD_ALIGN_PARAGRAPH.RIGHT
        p_date.add_run(f"\n")
        title.add_run(f"\n")
        p0_data=document.add_paragraph()
        p0_data.add_run(f"ID\tDoctor Name\tPatient Name\tIntervention\tBill\tDate")
        p0_data.add_run("\n")
        for item in data_list:
            p_data=document.add_paragraph()
            doc_id, doc_name, pat_name, interv, price, q_date = item
            p_data.add_run(f"{doc_id}\t{doc_name}\t{pat_name}\t{interv}\t{str(price)}\t{q_date}")
            p_data.add_run(f"\n")
            p_data.add_run(f"\n")
        p_data.add_run(f"Total: ").bold=True
        p_data.add_run(f"{str(total)} FCFA").bold=True
        p_data.add_run(f"\n")
        p_data.add_run(f"\t{word.upper()} FCFA").italic = True
        
        
    def generate_quote(self):
        doc = Document(os.path.join(os.getcwd(), resource_path('template.docx')))
        if self.var_doc_name.get()=='':
            messagebox.showerror("Error","Please select Doctor whose Quote you are Generating ",parent=self.root)
        else:
            try:
                self.get_quoteTable()
                current_month_name = str(time.strftime('%B'))
                self.add_quoting_data_to_document(doc,self.quote_data_list,self.var_doc_name.get(),self.var_date.get())
                doc.save(os.path.join(os.getcwd(), resource_path(f'Doctor_Quotes\\{self.var_doc_name.get()}_{current_month_name}_quote.docx')))  
                messagebox.showinfo("Success",f"Dr.{name.upper()}'s Quota for this Month Has Been Generated and saved",parent=self.root)
            except Exception as e:
                messagebox.showerror("Error",f"Error Due to: {str(e)}",parent=self.root)
                
    def view_quote(self):
        current_month_name = str(time.strftime('%B'))
        if self.var_doc_name.get()=='':
            messagebox.showerror("Error","Please Generate Quote first",parent=self.root)
        else: 
            file_path=os.path.join(os.getcwd(), resource_path(f'Doctor_Quotes\\{self.var_doc_name.get()}_{current_month_name}_quote.docx'))
            if not os.path.exists(file_path):
                messagebox.showerror("View Error","Generate Patient Proforma First before viewing",parent=self.root)
            else:
                os.startfile(os.path.join(os.getcwd(),resource_path(file_path)))         
       
    def print_quote(self):
        _this_month = str(time.strftime('%B'))
        file_path=os.path.join(os.getcwd(), resource_path(f'Doctor_Quotes\\{self.var_doc_name.get()}_{_this_month}_quote.docx'))
        if not os.path.exists(file_path):
            messagebox.showerror("Print error",f"The file {file_path} does not exist.Generate Doctor Quote First",parent=self.root)
        else:
            try:
                notification.notify(title="printer",message="Printing Process on-going!!!, please wait...",timeout=30)
                doc=docx.Document(file_path)
                printer_name=win32print.GetDefaultPrinter()
                win32api.ShellExecute(0,"print",file_path,'d:"%s"'%printer_name,".",0)
            except Exception as e:
                messagebox.Showerror("Printing Error",f"Error Due to: {str(e)}",parent=self.root)
            
    def clear_all(self):
        del self.quote_data_list[:]
        del self.quote_list[:]
        self.lbl_amnt.configure(text=f'Quotes(XAF)\n[0]')
        self.lbl_net_pay.configure(text=f'Net Quote(XAF)\n[0]')
        self.quoteTitle.configure(text=f"Quote \t Total Quotes: [0]")
        self.var_search.set("")
        self.show()
        self.show_quote()
        self.var_doc_id.set('')
        self.var_doc_name.set('')
        self.var_pat_name.set('')
        self.var_interv.set('')
        self.var_amt_paid.set('')
        self.var_date.set('')
    
    def update_date_time(self):
        time_ = time.strftime("%H:%M:%S")
        date_ = time.strftime("%d:%m:%Y")
        self.lbl_clock.configure(text=f"Welcome to SmileScribePro - Professional Patient Records Management System\t\t Date: {str(date_)}\t\t Time: {str(time_)}")
        self.lbl_clock.after(200, self.update_date_time)
        

    def exit(self):
        self.is_running = False
        self.root.destroy()
    
            
    
        
        
if __name__=="__main__":
    root=ctk.CTk()
    obj=quoteClass(root)
    root.mainloop()